"""
Fluxo HTTP e serviços de aplicação do app `licitacoes`.

Este módulo concentra:
- controle de acesso e orquestração de CRUD para termo/sessão/subsessão/item;
- regras de ordenação hierárquica dos nós do termo;
- duplicação estrutural de TR;
- importação estruturada por DOCX e por texto colado;
- exportação do termo para DOCX com layout institucional.

Integrações relevantes:
- ORM Django (consulta e escrita transacional em árvore recursiva);
- `python-docx` para leitura/escrita de documentos Word;
- templates do app para renderização de edição e visualização.
"""

from collections import defaultdict
from io import BytesIO
import json
import logging
import re

from django.contrib import messages
from django.contrib.auth.mixins import UserPassesTestMixin
from django.db import transaction
from django.db.models import Max
from django.http import HttpResponse
from django.http import Http404
from django.http import JsonResponse
from django.shortcuts import get_object_or_404, redirect
from django.urls import reverse, reverse_lazy
from django.utils.text import slugify
from django.utils.http import url_has_allowed_host_and_scheme
from django.views import View
from django.views.generic import CreateView, DeleteView, DetailView, FormView, ListView, TemplateView, UpdateView

from .forms import (
    ItemSessaoForm,
    EtpTicCreateForm,
    EtpTicSecaoForm,
    SessaoTermoForm,
    SubsessaoTermoForm,
    TermoReferenciaImportarForm,
    TabelaItemLinhaForm,
    TermoReferenciaDuplicarForm,
    TermoReferenciaForm,
)
from .models import EtpTic, ItemSessao, SessaoTermo, SubsessaoTermo, TabelaItemLinha, TermoReferencia

IMPORT_COMMENT_PREFIX = "NÃO SE ESQUEÇA DE RETIRAR ESSE COMENTÁRIO: "
# Modo estrito: registra linhas ambíguas no log para facilitar calibração do parser.
TEXT_IMPORT_STRICT_MODE = True
TEXT_IMPORT_MAX_HEADING_CHARS = 120
TEXT_IMPORT_MAX_HEADING_WORDS = 12

logger = logging.getLogger(__name__)


SEDS_REPLACEMENT = "Secretaria de Desenvolvimento Social"


def _expand_seds_references(text: str) -> str:
    """Expande a sigla institucional nas visualizações públicas do TR."""

    return re.sub(r"\bSEDS\b", SEDS_REPLACEMENT, text or "")


ETP_TIC_SECOES = [
    {
        "numero": 1,
        "titulo": "Informações Básicas",
        "descricao": "",
        "campos": ["numero_processo_servico", "titulo"],
    },
    {
        "numero": 2,
        "titulo": "Descrição da Necessidade",
        "descricao": (
            "Descrição da necessidade para o órgão, entidade ou determinado público-alvo que a "
            "contratação se propõe a atender. A descrição da necessidade pode ser realizada por meio "
            "da declaração da situação-problema a ser enfrentada ou da oportunidade de negócio a ser "
            "alcançada, incluindo a discriminação do valor a ser entregue à organização ou ao público-alvo."
        ),
        "campos": ["descricao_necessidade"],
    },
    {
        "numero": 3,
        "titulo": "Área Requisitante",
        "descricao": "",
        "campos": ["area_requisitante", "responsavel_area"],
    },
    {
        "numero": 4,
        "titulo": "Necessidades de Negócio",
        "descricao": (
            "Conforme o inciso I do art. 11 da IN SGD/ME nº 1, de 2019, o órgão demandante deve "
            "definir e especificar as necessidades de negócio que a contratação visa atender."
        ),
        "campos": ["necessidades_negocio"],
    },
    {
        "numero": 5,
        "titulo": "Necessidades Tecnológicas",
        "descricao": (
            "Conforme o inciso I do art. 11 da IN SGD/ME nº 1, de 2019, o órgão demandante deve "
            "definir e especificar as necessidades tecnológicas."
        ),
        "campos": ["necessidades_tecnologicas"],
    },
    {
        "numero": 6,
        "titulo": "Demais Requisitos Necessários e Suficientes à Escolha da Solução de TIC",
        "descricao": (
            "Devem ser descritos todos os requisitos indispensáveis ao atendimento da necessidade de "
            "negócio, garantindo-se a economicidade da contratação."
        ),
        "campos": ["demais_requisitos"],
    },
    {
        "numero": 7,
        "titulo": "Estimativa da Demanda - Quantidade de Bens e Serviços",
        "descricao": (
            "Conforme o inciso I do art. 11 da IN SGD/ME nº 1, de 2019, o órgão demandante deve "
            "descrever de forma detalhada, motivada e justificada o quantitativo de bens e serviços."
        ),
        "campos": ["estimativa_demanda"],
    },
    {
        "numero": 8,
        "titulo": "Levantamento de Soluções",
        "descricao": (
            "Deve-se realizar um levantamento de soluções disponíveis que possam atender à necessidade "
            "da contratação para o órgão ou entidade."
        ),
        "campos": ["levantamento_solucoes"],
    },
    {
        "numero": 9,
        "titulo": "Análise Comparativa de Soluções",
        "descricao": (
            "Consiste em uma análise crítica entre as diferentes soluções, considerando o aspecto "
            "econômico (TCO) e os aspectos qualitativos."
        ),
        "campos": ["analise_comparativa_solucoes"],
    },
    {
        "numero": 10,
        "titulo": "Registro de Soluções Consideradas Inviáveis",
        "descricao": (
            "Conforme o § 1º do art. 11, as soluções identificadas e consideradas inviáveis deverão "
            "ser registradas no Estudo Técnico Preliminar da Contratação."
        ),
        "campos": ["solucoes_inviaveis"],
    },
    {
        "numero": 11,
        "titulo": "Análise Comparativa de Custos (TCO)",
        "descricao": (
            "Conforme o inciso III do art. 11, deve-se proceder à comparação de custos totais de "
            "propriedade para as soluções técnica e funcionalmente viáveis."
        ),
        "campos": ["analise_comparativa_custos_tco"],
    },
    {
        "numero": 12,
        "titulo": "Descrição da Solução de TIC a Ser Contratada",
        "descricao": "Deve-se identificar a solução escolhida que será objeto da contratação.",
        "campos": ["descricao_solucao_tic"],
    },
    {
        "numero": 13,
        "titulo": "Estimativa de Custo Total da Contratação",
        "descricao": "Registro da estimativa do custo da contratação, considerando a solução escolhida.",
        "campos": ["estimativa_custo_valor", "estimativa_custo_texto"],
    },
    {
        "numero": 14,
        "titulo": "Justificativa Técnica da Escolha da Solução",
        "descricao": (
            "Descrever as razões técnicas que motivaram a escolha da solução, considerando os "
            "aspectos qualitativos, incluindo riscos e benefícios esperados."
        ),
        "campos": ["justificativa_tecnica"],
    },
    {
        "numero": 15,
        "titulo": "Justificativa Econômica da Escolha da Solução",
        "descricao": (
            "Descrever as razões, em termos quantitativos, que motivaram a escolha da solução, "
            "enfatizando os aspectos de economicidade."
        ),
        "campos": ["justificativa_economica"],
    },
    {
        "numero": 16,
        "titulo": "Benefícios a Serem Alcançados com a Contratação",
        "descricao": (
            "Identificação dos resultados a serem alcançados com a contratação da solução de TIC."
        ),
        "campos": ["beneficios_contratacao"],
    },
    {
        "numero": 17,
        "titulo": "Providências a Serem Adotadas",
        "descricao": (
            "Informar, se houver, todas as providências e as necessidades de adequação do ambiente."
        ),
        "campos": ["providencias_adotadas"],
    },
    {
        "numero": 18,
        "titulo": "Declaração de Viabilidade",
        "descricao": "",
        "campos": ["declaracao_viabilidade", "justificativa_viabilidade"],
    },
]

ETP_TIC_SECOES_MAP = {secao["numero"]: secao for secao in ETP_TIC_SECOES}


def _etp_split_paragraphs(texto: str) -> list[str]:
    blocos = re.split(r"\n\s*\n", (texto or "").strip())
    return [bloco.strip() for bloco in blocos if bloco and bloco.strip()]


def _etp_is_secao_preenchida(etp: EtpTic, secao_numero: int) -> bool:
    secao = ETP_TIC_SECOES_MAP[secao_numero]
    for campo in secao["campos"]:
        valor = getattr(etp, campo, None)
        if valor is None:
            continue
        if isinstance(valor, str) and valor.strip():
            return True
        if not isinstance(valor, str) and valor not in ("", None):
            return True
    return False


def _etp_status_por_secao(etp: EtpTic) -> dict[int, str]:
    status = {}
    for secao in ETP_TIC_SECOES:
        numero = secao["numero"]
        preenchida = _etp_is_secao_preenchida(etp, numero)
        if numero < etp.secao_atual:
            status[numero] = "concluido" if preenchida else "nao_iniciado"
        elif numero == etp.secao_atual:
            status[numero] = "em_andamento" if preenchida else "nao_iniciado"
        else:
            status[numero] = "nao_iniciado"
    return status


def _etp_render_sections(etp: EtpTic) -> list[dict]:
    secoes_render = []
    for secao in ETP_TIC_SECOES:
        numero = secao["numero"]
        entradas = []
        if numero == 13:
            if etp.estimativa_custo_valor is not None:
                entradas.append(f"Valor (R$): {etp.estimativa_custo_valor}")
            paragrafos = _etp_split_paragraphs(etp.estimativa_custo_texto)
            for idx, paragrafo in enumerate(paragrafos, start=1):
                entradas.append(f"{numero}.{idx}. {paragrafo}")
        elif numero == 1:
            if etp.numero_processo_servico:
                entradas.append(f"{numero}.1. Número do processo de serviço: {etp.numero_processo_servico}")
        elif numero == 3:
            if etp.area_requisitante:
                entradas.append(f"{numero}.1. Área requisitante: {etp.area_requisitante}")
            if etp.responsavel_area:
                entradas.append(f"{numero}.2. Responsável: {etp.responsavel_area}")
        elif numero == 18:
            if etp.declaracao_viabilidade:
                entradas.append(f"{numero}.1. {etp.declaracao_viabilidade.strip()}")
            for idx, paragrafo in enumerate(_etp_split_paragraphs(etp.justificativa_viabilidade), start=1):
                entradas.append(f"{numero}.{idx + 1}. {paragrafo}")
        else:
            campo = secao["campos"][0]
            paragrafos = _etp_split_paragraphs(getattr(etp, campo, ""))
            for idx, paragrafo in enumerate(paragrafos, start=1):
                entradas.append(f"{numero}.{idx}. {paragrafo}")
        secoes_render.append(
            {
                "numero": numero,
                "titulo": secao["titulo"],
                "descricao": secao["descricao"],
                "entradas": entradas,
            }
        )
    return secoes_render


def _can_access_licitacoes(user) -> bool:
    """Função `_can_access_licitacoes` no fluxo do app `licitacoes`.

    Objetivo pedagógico:
    - Executa uma etapa específica de validação, transformação ou orquestração.
    - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
    - Parâmetros: `user`.
    """

    if not user.is_authenticated:
        return False
    return (
        user.is_staff
        or user.has_perm("licitacoes.view_termoreferencia")
        or user.has_perm("licitacoes.add_termoreferencia")
        or user.has_perm("licitacoes.change_termoreferencia")
        or user.has_perm("licitacoes.delete_termoreferencia")
        or user.has_perm("licitacoes.view_sessaotermo")
        or user.has_perm("licitacoes.add_sessaotermo")
        or user.has_perm("licitacoes.change_sessaotermo")
        or user.has_perm("licitacoes.delete_sessaotermo")
        or user.has_perm("licitacoes.view_subsessaotermo")
        or user.has_perm("licitacoes.add_subsessaotermo")
        or user.has_perm("licitacoes.change_subsessaotermo")
        or user.has_perm("licitacoes.delete_subsessaotermo")
        or user.has_perm("licitacoes.view_itemsessao")
        or user.has_perm("licitacoes.add_itemsessao")
        or user.has_perm("licitacoes.change_itemsessao")
        or user.has_perm("licitacoes.delete_itemsessao")
        or user.has_perm("licitacoes.view_tabelaitemlinha")
        or user.has_perm("licitacoes.add_tabelaitemlinha")
        or user.has_perm("licitacoes.change_tabelaitemlinha")
        or user.has_perm("licitacoes.delete_tabelaitemlinha")
    )


def _normalize_sessoes(termo: TermoReferencia) -> None:
    """Função `_normalize_sessoes` no fluxo do app `licitacoes`.

    Objetivo pedagógico:
    - Executa uma etapa específica de validação, transformação ou orquestração.
    - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
    - Parâmetros: `termo`.
    """

    for idx, sessao in enumerate(termo.sessoes.order_by("ordem", "id"), start=1):
        if sessao.ordem != idx:
            sessao.ordem = idx
            sessao.save(update_fields=["ordem"])


def _normalize_subsessoes(sessao: SessaoTermo) -> None:
    """Função `_normalize_subsessoes` no fluxo do app `licitacoes`.

    Objetivo pedagógico:
    - Executa uma etapa específica de validação, transformação ou orquestração.
    - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
    - Parâmetros: `sessao`.
    """

    for idx, subsessao in enumerate(sessao.subsessoes.order_by("ordem", "id"), start=1):
        if subsessao.ordem != idx:
            subsessao.ordem = idx
            subsessao.save(update_fields=["ordem"])


def _normalize_items(sessao: SessaoTermo, parent_id: int | None) -> None:
    """Função `_normalize_items` no fluxo do app `licitacoes`.

    Objetivo pedagógico:
    - Executa uma etapa específica de validação, transformação ou orquestração.
    - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
    - Parâmetros: `sessao`, `parent_id`.
    """

    qs = sessao.itens.filter(parent_id=parent_id).order_by("ordem", "id")
    for idx, item in enumerate(qs, start=1):
        if item.ordem != idx:
            item.ordem = idx
            item.save(update_fields=["ordem"])


def _next_ordem_sessao(termo: TermoReferencia) -> int:
    """Função `_next_ordem_sessao` no fluxo do app `licitacoes`.

    Objetivo pedagógico:
    - Executa uma etapa específica de validação, transformação ou orquestração.
    - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
    - Parâmetros: `termo`.
    """

    current_max = termo.sessoes.aggregate(max_ordem=Max("ordem"))["max_ordem"]
    return (current_max or 0) + 1


def _next_ordem_subsessao(sessao: SessaoTermo) -> int:
    """Função `_next_ordem_subsessao` no fluxo do app `licitacoes`.

    Objetivo pedagógico:
    - Executa uma etapa específica de validação, transformação ou orquestração.
    - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
    - Parâmetros: `sessao`.
    """

    current_max = sessao.subsessoes.aggregate(max_ordem=Max("ordem"))["max_ordem"]
    return (current_max or 0) + 1


def _next_ordem_item(sessao: SessaoTermo, parent_id: int | None) -> int:
    """Função `_next_ordem_item` no fluxo do app `licitacoes`.

    Objetivo pedagógico:
    - Executa uma etapa específica de validação, transformação ou orquestração.
    - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
    - Parâmetros: `sessao`, `parent_id`.
    """

    current_max = sessao.itens.filter(parent_id=parent_id).aggregate(max_ordem=Max("ordem"))["max_ordem"]
    return (current_max or 0) + 1


def _ordered_item_ids(sessao: SessaoTermo, parent_id: int | None, exclude_id: int | None = None) -> list[int]:
    """Função `_ordered_item_ids` no fluxo do app `licitacoes`.

    Objetivo pedagógico:
    - Executa uma etapa específica de validação, transformação ou orquestração.
    - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
    - Parâmetros: `sessao`, `parent_id`, `exclude_id`.
    """

    ids = list(
        sessao.itens.filter(parent_id=parent_id)
        .order_by("ordem", "id")
        .values_list("id", flat=True)
    )
    if exclude_id is None:
        return ids
    return [item_id for item_id in ids if item_id != exclude_id]


def _apply_group_order(ordered_ids: list[int]) -> None:
    """Função `_apply_group_order` no fluxo do app `licitacoes`.

    Objetivo pedagógico:
    - Executa uma etapa específica de validação, transformação ou orquestração.
    - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
    - Parâmetros: `ordered_ids`.
    """

    for idx, item_id in enumerate(ordered_ids, start=1):
        ItemSessao.objects.filter(pk=item_id).update(ordem=idx)


def _next_ordem_tabela(item: ItemSessao) -> int:
    """Função `_next_ordem_tabela` no fluxo do app `licitacoes`.

    Objetivo pedagógico:
    - Executa uma etapa específica de validação, transformação ou orquestração.
    - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
    - Parâmetros: `item`.
    """

    current_max = item.tabela_linhas.aggregate(max_ordem=Max("ordem"))["max_ordem"]
    return (current_max or 0) + 1


def _normalize_tabela_linhas(item: ItemSessao) -> None:
    """Função `_normalize_tabela_linhas` no fluxo do app `licitacoes`.

    Objetivo pedagógico:
    - Executa uma etapa específica de validação, transformação ou orquestração.
    - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
    - Parâmetros: `item`.
    """

    for idx, linha in enumerate(item.tabela_linhas.order_by("ordem", "id"), start=1):
        if linha.ordem != idx:
            linha.ordem = idx
            linha.save(update_fields=["ordem"])


def _item_ordem_exibicao(item: ItemSessao, siblings: list[ItemSessao]) -> int:
    """Função `_item_ordem_exibicao` no fluxo do app `licitacoes`.

    Objetivo pedagógico:
    - Executa uma etapa específica de validação, transformação ou orquestração.
    - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
    - Parâmetros: `item`, `siblings`.
    """

    if item.enum_tipo in {ItemSessao.EnumTipo.INCISO, ItemSessao.EnumTipo.ALINEA}:
        mesma_familia = [s for s in siblings if s.enum_tipo == item.enum_tipo]
    else:
        mesma_familia = [s for s in siblings if s.enum_tipo in {"", None}]
    for idx, sibling in enumerate(mesma_familia, start=1):
        if sibling.id == item.id:
            return idx
    return 1


def _item_indice(item: ItemSessao) -> str:
    """Função `_item_indice` no fluxo do app `licitacoes`.

    Objetivo pedagógico:
    - Executa uma etapa específica de validação, transformação ou orquestração.
    - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
    - Parâmetros: `item`.
    """

    niveis = [str(item.sessao.ordem)]
    hierarquia: list[ItemSessao] = []
    atual = item
    while atual:
        hierarquia.append(atual)
        atual = atual.parent
    for atual in reversed(hierarquia):
        siblings = list(
            atual.sessao.itens.filter(parent_id=atual.parent_id).order_by("ordem", "id")
        )
        niveis.append(str(_item_ordem_exibicao(atual, siblings)))
    return ".".join(niveis)


def _can_have_tabela_itens(item: ItemSessao) -> bool:
    """Função `_can_have_tabela_itens` no fluxo do app `licitacoes`.

    Objetivo pedagógico:
    - Executa uma etapa específica de validação, transformação ou orquestração.
    - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
    - Parâmetros: `item`.
    """

    return _item_indice(item) == "1.1"


def _quantidade_text(value) -> str:
    """Função `_quantidade_text` no fluxo do app `licitacoes`.

    Objetivo pedagógico:
    - Executa uma etapa específica de validação, transformação ou orquestração.
    - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
    - Parâmetros: `value`.
    """

    if value is None:
        return ""
    as_text = f"{value}".replace(".", ",")
    if as_text.endswith(",00"):
        return as_text[:-3]
    return as_text


def _int_to_roman(num: int) -> str:
    """Função `_int_to_roman` no fluxo do app `licitacoes`.

    Objetivo pedagógico:
    - Executa uma etapa específica de validação, transformação ou orquestração.
    - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
    - Parâmetros: `num`.
    """

    pairs = [
        (1000, "M"), (900, "CM"), (500, "D"), (400, "CD"),
        (100, "C"), (90, "XC"), (50, "L"), (40, "XL"),
        (10, "X"), (9, "IX"), (5, "V"), (4, "IV"), (1, "I"),
    ]
    out = []
    n = num
    for val, sym in pairs:
        while n >= val:
            out.append(sym)
            n -= val
    return "".join(out)


def _enum_prefix(item: ItemSessao, siblings: list[ItemSessao]) -> str:
    """Função `_enum_prefix` no fluxo do app `licitacoes`.

    Objetivo pedagógico:
    - Executa uma etapa específica de validação, transformação ou orquestração.
    - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
    - Parâmetros: `item`, `siblings`.
    """

    if item.enum_tipo == ItemSessao.EnumTipo.INCISO:
        ordem = _item_ordem_exibicao(item, siblings)
        return f"{_int_to_roman(ordem)})"
    if item.enum_tipo == ItemSessao.EnumTipo.ALINEA:
        ordem = _item_ordem_exibicao(item, siblings)
        letter = chr(ord("a") + max(0, ordem - 1))
        return f"{letter})"
    return ""


def _display_item_text(item: ItemSessao) -> str:
    """Função `_display_item_text` no fluxo do app `licitacoes`.

    Objetivo pedagógico:
    - Executa uma etapa específica de validação, transformação ou orquestração.
    - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
    - Parâmetros: `item`.
    """

    texto = (item.texto or "").lstrip()
    if item.enum_tipo in {ItemSessao.EnumTipo.ALINEA, ItemSessao.EnumTipo.INCISO}:
        m = re.match(r"^\s*([IVXLCDM])\)\s+([a-z].*)$", texto)
        if m:
            return _expand_seds_references(f"{m.group(1)}{m.group(2)}")
    return _expand_seds_references(texto)


def _safe_next_url(request, default_url: str) -> str:
    """Função `_safe_next_url` no fluxo do app `licitacoes`.

    Objetivo pedagógico:
    - Executa uma etapa específica de validação, transformação ou orquestração.
    - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
    - Parâmetros: `request`, `default_url`.
    """

    next_url = request.POST.get("next") or request.GET.get("next")
    if next_url and url_has_allowed_host_and_scheme(
        url=next_url,
        allowed_hosts={request.get_host()},
        require_https=request.is_secure(),
    ):
        return next_url
    return default_url


def _build_item_rows(sessao: SessaoTermo, sessao_indice: int) -> list[dict]:
    """Função `_build_item_rows` no fluxo do app `licitacoes`.

    Objetivo pedagógico:
    - Executa uma etapa específica de validação, transformação ou orquestração.
    - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
    - Parâmetros: `sessao`, `sessao_indice`.
    """

    # Carrega itens da sessão de uma vez e com relações necessárias para evitar N+1 no template.
    itens = list(sessao.itens.select_related("parent", "subsessao").order_by("parent_id", "ordem", "id"))
    by_parent: dict[int | None, list[ItemSessao]] = defaultdict(list)
    for item in itens:
        by_parent[item.parent_id].append(item)

    rows: list[dict] = []

    def walk(parent_id: int | None, parent_index: str, depth: int) -> None:
        """Função `_build_item_rows.walk` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `parent_id`, `parent_index`, `depth`.
        """

        # O índice de exibição é calculado por grupo de irmãos, preservando semântica de enumeração.
        siblings = by_parent.get(parent_id, [])
        for item in siblings:
            indice = f"{parent_index}.{_item_ordem_exibicao(item, siblings)}"
            rows.append(
                {
                    "item": item,
                    "indice": indice,
                    "enum_prefix": _enum_prefix(item, siblings),
                    "display_text": _display_item_text(item),
                    "depth": depth,
                    "indent_px": depth * 24,
                }
            )
            walk(item.id, indice, depth + 1)

    walk(None, str(sessao_indice), 0)
    return rows


def _duplicar_estrutura_termo(origem: TermoReferencia, novo_apelido: str) -> TermoReferencia:
    """Função `_duplicar_estrutura_termo` no fluxo do app `licitacoes`.

    Objetivo pedagógico:
    - Executa uma etapa específica de validação, transformação ou orquestração.
    - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
    - Parâmetros: `origem`, `novo_apelido`.
    """

    # A duplicação é transacional: ou replica toda a árvore do termo, ou não grava nada.
    with transaction.atomic():
        novo_termo = TermoReferencia.objects.create(
            apelido=novo_apelido,
            processo_sei=origem.processo_sei,
            link_processo_sei=origem.link_processo_sei,
        )

        sessoes_map: dict[int, SessaoTermo] = {}
        # Primeiro pass: cria sessões para manter mapa de IDs antigos -> novos.
        for sessao_origem in origem.sessoes.order_by("ordem", "id"):
            sessoes_map[sessao_origem.id] = SessaoTermo.objects.create(
                termo=novo_termo,
                titulo=sessao_origem.titulo,
                ordem=sessao_origem.ordem,
            )

        for sessao_origem in origem.sessoes.order_by("ordem", "id"):
            sessao_nova = sessoes_map[sessao_origem.id]
            subsessoes_map: dict[int, SubsessaoTermo] = {}
            for sub_origem in sessao_origem.subsessoes.order_by("ordem", "id"):
                subsessoes_map[sub_origem.id] = SubsessaoTermo.objects.create(
                    sessao=sessao_nova,
                    titulo=sub_origem.titulo,
                    ordem=sub_origem.ordem,
                )
            filhos_por_parent: dict[int | None, list[ItemSessao]] = defaultdict(list)
            # Pré-agrupa itens por pai para permitir clonagem recursiva preservando a hierarquia.
            for item in sessao_origem.itens.order_by("ordem", "id"):
                filhos_por_parent[item.parent_id].append(item)

            def clonar_itens(parent_origem_id: int | None, parent_novo: ItemSessao | None) -> None:
                """Função `_duplicar_estrutura_termo.clonar_itens` no fluxo do app `licitacoes`.

                Objetivo pedagógico:
                - Executa uma etapa específica de validação, transformação ou orquestração.
                - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
                - Parâmetros: `parent_origem_id`, `parent_novo`.
                """

                for item_origem in filhos_por_parent.get(parent_origem_id, []):
                    item_novo = ItemSessao.objects.create(
                        sessao=sessao_nova,
                        subsessao=subsessoes_map.get(item_origem.subsessao_id),
                        parent=parent_novo,
                        enum_tipo=item_origem.enum_tipo,
                        texto=item_origem.texto,
                        ordem=item_origem.ordem,
                    )
                    # Replica também a tabela vinculada ao item, mantendo ordem e conteúdo.
                    for linha in item_origem.tabela_linhas.order_by("ordem", "id"):
                        TabelaItemLinha.objects.create(
                            item=item_novo,
                            ordem=linha.ordem,
                            descricao=linha.descricao,
                            catmat_catser=linha.catmat_catser,
                            siafisico=linha.siafisico,
                            unidade_fornecimento=linha.unidade_fornecimento,
                            quantidade=linha.quantidade,
                        )
                    clonar_itens(item_origem.id, item_novo)

            clonar_itens(None, None)

        return novo_termo


def _normalize_docx_style_name(style_name: str) -> str:
    """Função `_normalize_docx_style_name` no fluxo do app `licitacoes`.

    Objetivo pedagógico:
    - Executa uma etapa específica de validação, transformação ou orquestração.
    - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
    - Parâmetros: `style_name`.
    """

    translation = str.maketrans("áàâãéêíóôõúç", "aaaaeeiooouc")
    normalized = (style_name or "").lower().translate(translation)
    return re.sub(r"[^a-z0-9]+", "", normalized)


def _docx_level_from_style(style_name: str) -> int | None:
    """Função `_docx_level_from_style` no fluxo do app `licitacoes`.

    Objetivo pedagógico:
    - Executa uma etapa específica de validação, transformação ou orquestração.
    - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
    - Parâmetros: `style_name`.
    """

    token = _normalize_docx_style_name(style_name)
    if re.search(r"(nivel|nvel)0?1", token):
        return 1
    if re.search(r"(nivel|nvel)0?2", token):
        return 2
    if re.search(r"(nivel|nvel)0?3", token):
        return 3
    if re.search(r"(nivel|nvel)0?4", token):
        return 4
    return None


def _strip_docx_heading_prefix(texto: str) -> str:
    """Função `_strip_docx_heading_prefix` no fluxo do app `licitacoes`.

    Objetivo pedagógico:
    - Executa uma etapa específica de validação, transformação ou orquestração.
    - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
    - Parâmetros: `texto`.
    """

    sem_numeracao = re.sub(r"^\s*\d+(?:\.\d+)*(?:\.)?\s*[-–—:]?\s*", "", texto.strip())
    return sem_numeracao.strip()


def _truncate_to_model_field(value: str, max_length: int) -> str:
    """Função `_truncate_to_model_field` no fluxo do app `licitacoes`.

    Objetivo pedagógico:
    - Executa uma etapa específica de validação, transformação ou orquestração.
    - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
    - Parâmetros: `value`, `max_length`.
    """

    cleaned = (value or "").strip()
    if len(cleaned) <= max_length:
        return cleaned
    return cleaned[: max_length - 3].rstrip() + "..."


def _is_alternative_or_comment_line(texto: str) -> bool:
    """Função `_is_alternative_or_comment_line` no fluxo do app `licitacoes`.

    Objetivo pedagógico:
    - Executa uma etapa específica de validação, transformação ou orquestração.
    - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
    - Parâmetros: `texto`.
    """

    token = texto.strip().lower()
    if token == "ou":
        return True
    return "alternativa" in token or "coment" in token


def _normalize_import_line(texto: str) -> str:
    """Função `_normalize_import_line` no fluxo do app `licitacoes`.

    Objetivo pedagógico:
    - Executa uma etapa específica de validação, transformação ou orquestração.
    - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
    - Parâmetros: `texto`.
    """

    cleaned = (texto or "").replace("\u00a0", " ").replace("\u200b", " ")
    cleaned = re.sub(r"[ \t]+", " ", cleaned)
    return cleaned.strip()


def _looks_like_plain_heading(texto: str) -> bool:
    """Função `_looks_like_plain_heading` no fluxo do app `licitacoes`.

    Objetivo pedagógico:
    - Executa uma etapa específica de validação, transformação ou orquestração.
    - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
    - Parâmetros: `texto`.
    """

    if not texto:
        return False
    if len(texto) > TEXT_IMPORT_MAX_HEADING_CHARS:
        return False
    if len(texto.split()) > TEXT_IMPORT_MAX_HEADING_WORDS:
        return False
    if texto.endswith((".", ";", ":", "?", "!")):
        return False
    return True


def _tokenize_import_line(texto: str) -> tuple[str, dict[str, str]]:
    """Função `_tokenize_import_line` no fluxo do app `licitacoes`.

    Objetivo pedagógico:
    - Executa uma etapa específica de validação, transformação ou orquestração.
    - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
    - Parâmetros: `texto`.
    """

    section_match = re.match(r"^(\d+)\.\s+(.+)$", texto)
    if section_match:
        return "SECTION_NUM", {"numero": section_match.group(1), "texto": section_match.group(2)}

    numeric_match = re.match(r"^(\d+(?:\.\d+)+)\.\s+(.+)$", texto)
    if numeric_match:
        return "NUMERIC_ITEM", {"numero": numeric_match.group(1), "texto": numeric_match.group(2)}

    roman_match = re.match(r"^([IVXLCDM]+)\)\s+(.+)$", texto, flags=re.IGNORECASE)
    if roman_match:
        return "ROMAN_INCISO", {"enum": roman_match.group(1), "texto": roman_match.group(2)}

    alpha_match = re.match(r"^([a-z])\)\s+(.+)$", texto, flags=re.IGNORECASE)
    if alpha_match:
        return "ALPHA_ALINEA", {"enum": alpha_match.group(1), "texto": alpha_match.group(2)}

    if re.match(r"^OU$", texto, flags=re.IGNORECASE):
        return "OU_LINE", {"texto": texto}

    if re.match(r"^\[.*\]$", texto):
        return "BRACKET_NOTE", {"texto": texto}

    if _looks_like_plain_heading(texto):
        return "PLAIN_HEADING", {"texto": texto}

    return "PLAIN_TEXT", {"texto": texto}


def _apply_import_comment_prefix(texto: str) -> str:
    """Função `_apply_import_comment_prefix` no fluxo do app `licitacoes`.

    Objetivo pedagógico:
    - Executa uma etapa específica de validação, transformação ou orquestração.
    - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
    - Parâmetros: `texto`.
    """

    valor = texto.strip()
    if not valor:
        return valor
    if valor.startswith(IMPORT_COMMENT_PREFIX):
        return valor
    return f"{IMPORT_COMMENT_PREFIX}{valor}"


def _extract_item_text_and_enum(raw_text: str) -> tuple[str, str]:
    """Função `_extract_item_text_and_enum` no fluxo do app `licitacoes`.

    Objetivo pedagógico:
    - Executa uma etapa específica de validação, transformação ou orquestração.
    - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
    - Parâmetros: `raw_text`.
    """

    texto = _strip_docx_heading_prefix(raw_text)
    roman_match = re.match(r"^\s*([IVXLCDM]+)\)\s*(.+)$", texto, flags=re.IGNORECASE)
    if roman_match:
        return roman_match.group(2).strip(), ItemSessao.EnumTipo.INCISO
    alpha_match = re.match(r"^\s*([a-z])\)\s*(.+)$", texto, flags=re.IGNORECASE)
    if alpha_match:
        return alpha_match.group(2).strip(), ItemSessao.EnumTipo.ALINEA
    return texto, ItemSessao.EnumTipo.NENHUM


def _importar_termo_texto(texto_bruto: str, apelido: str) -> tuple[TermoReferencia, dict[str, int]]:
    """Função `_importar_termo_texto` no fluxo do app `licitacoes`.

    Objetivo pedagógico:
    - Executa uma etapa específica de validação, transformação ou orquestração.
    - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
    - Parâmetros: `texto_bruto`, `apelido`.
    """

    max_titulo_sessao = SessaoTermo._meta.get_field("titulo").max_length
    max_titulo_subsessao = SubsessaoTermo._meta.get_field("titulo").max_length
    # Normaliza e remove linhas vazias para reduzir ruído do texto colado do Word.
    linhas = [_normalize_import_line(ln) for ln in (texto_bruto or "").splitlines()]
    linhas = [ln for ln in linhas if ln]

    with transaction.atomic():
        termo = TermoReferencia.objects.create(
            apelido=apelido,
            processo_sei="IMPORTADO TEXTO",
            link_processo_sei="",
        )
        # Contadores retornados à UI para feedback operacional pós-importação.
        counters = {
            "sessoes": 0,
            "subsessoes": 0,
            "itens": 0,
            "incisos": 0,
            "alineas": 0,
            "comentarios": 0,
            "nao_classificadas": 0,
        }
        parse_report: list[dict[str, str]] = []

        def register_report(line_no: int, token: str, action: str, note: str) -> None:
            # Relatório técnico de fallback/ambiguidade (não exibido ao usuário final).
            parse_report.append(
                {
                    "linha": str(line_no),
                    "token": token,
                    "acao": action,
                    "detalhe": note,
                }
            )

        sessao_atual: SessaoTermo | None = None
        sessao_numero_atual: str | None = None
        subsessao_atual: SubsessaoTermo | None = None
        itens_por_numeracao: dict[str, ItemSessao] = {}
        item_numerico_corrente: ItemSessao | None = None
        inciso_corrente: ItemSessao | None = None
        ultimo_item_valido: ItemSessao | None = None
        ultimo_item_estrutural: ItemSessao | None = None

        def ensure_sessao(numero: str, line_no: int) -> None:
            # Garante âncora de sessão para itens numerados fora de contexto explícito.
            nonlocal sessao_atual, sessao_numero_atual, subsessao_atual
            nonlocal itens_por_numeracao, item_numerico_corrente, inciso_corrente, ultimo_item_valido, ultimo_item_estrutural

            if sessao_atual is not None and sessao_numero_atual == numero:
                return
            sessao_atual = SessaoTermo.objects.create(
                termo=termo,
                titulo=_truncate_to_model_field(f"Seção {numero}", max_titulo_sessao),
                ordem=_next_ordem_sessao(termo),
            )
            counters["sessoes"] += 1
            sessao_numero_atual = numero
            subsessao_atual = None
            itens_por_numeracao = {}
            item_numerico_corrente = None
            inciso_corrente = None
            ultimo_item_valido = None
            ultimo_item_estrutural = None
            if TEXT_IMPORT_STRICT_MODE:
                register_report(line_no, "IMPLICIT_SECTION", "create_section", "Sessão implícita por item numerado")

        for idx, linha in enumerate(linhas, start=1):
            # Tokenização determinística: cada linha cai em uma categoria de parsing.
            token, data = _tokenize_import_line(linha)

            if token == "SECTION_NUM":
                sessao_atual = SessaoTermo.objects.create(
                    termo=termo,
                    titulo=_truncate_to_model_field(data["texto"], max_titulo_sessao),
                    ordem=_next_ordem_sessao(termo),
                )
                counters["sessoes"] += 1
                sessao_numero_atual = data["numero"]
                subsessao_atual = None
                itens_por_numeracao = {}
                item_numerico_corrente = None
                inciso_corrente = None
                ultimo_item_valido = None
                ultimo_item_estrutural = None
                continue

            if token == "NUMERIC_ITEM":
                numero = data["numero"]
                texto_item = data["texto"]
                partes = numero.split(".")
                if sessao_atual is None or sessao_numero_atual != partes[0]:
                    ensure_sessao(partes[0], idx)
                parent = None
                # Para níveis 1.1.1+, vincula no ancestral imediato (1.1, 1.1.1, etc.).
                if len(partes) > 2:
                    parent_key = ".".join(partes[:-1])
                    parent = itens_por_numeracao.get(parent_key)
                    if parent is None and TEXT_IMPORT_STRICT_MODE:
                        counters["nao_classificadas"] += 1
                        register_report(idx, token, "fallback_parent_none", f"Pai {parent_key} não encontrado")
                novo_item = ItemSessao.objects.create(
                    sessao=sessao_atual,
                    subsessao=subsessao_atual,
                    parent=parent,
                    enum_tipo=ItemSessao.EnumTipo.NENHUM,
                    texto=texto_item,
                    ordem=_next_ordem_item(sessao_atual, parent.id if parent else None),
                )
                counters["itens"] += 1
                itens_por_numeracao[numero] = novo_item
                item_numerico_corrente = novo_item
                inciso_corrente = None
                ultimo_item_valido = novo_item
                ultimo_item_estrutural = novo_item
                continue

            if token == "PLAIN_HEADING":
                if sessao_atual is None:
                    sessao_atual = SessaoTermo.objects.create(
                        termo=termo,
                        titulo=_truncate_to_model_field("Conteudo importado", max_titulo_sessao),
                        ordem=_next_ordem_sessao(termo),
                    )
                    counters["sessoes"] += 1
                    sessao_numero_atual = None
                    if TEXT_IMPORT_STRICT_MODE:
                        register_report(idx, token, "create_default_section", "Título sem sessão numerada prévia")
                subsessao_atual = SubsessaoTermo.objects.create(
                    sessao=sessao_atual,
                    titulo=_truncate_to_model_field(data["texto"], max_titulo_subsessao),
                    ordem=_next_ordem_subsessao(sessao_atual),
                )
                counters["subsessoes"] += 1
                item_numerico_corrente = None
                inciso_corrente = None
                continue

            if token == "ROMAN_INCISO":
                if sessao_atual is None:
                    ensure_sessao("0", idx)
                parent = item_numerico_corrente or ultimo_item_valido
                if parent is None and TEXT_IMPORT_STRICT_MODE:
                    counters["nao_classificadas"] += 1
                    register_report(idx, token, "fallback_root", "Inciso sem item base; criado como raiz")
                novo_item = ItemSessao.objects.create(
                    sessao=sessao_atual,
                    subsessao=subsessao_atual,
                    parent=parent,
                    enum_tipo=ItemSessao.EnumTipo.INCISO if parent else ItemSessao.EnumTipo.NENHUM,
                    texto=data["texto"],
                    ordem=_next_ordem_item(sessao_atual, parent.id if parent else None),
                )
                counters["itens"] += 1
                counters["incisos"] += 1
                inciso_corrente = novo_item
                ultimo_item_valido = novo_item
                ultimo_item_estrutural = novo_item
                continue

            if token == "ALPHA_ALINEA":
                if sessao_atual is None:
                    ensure_sessao("0", idx)
                parent = inciso_corrente or item_numerico_corrente or ultimo_item_valido
                if parent is None and TEXT_IMPORT_STRICT_MODE:
                    counters["nao_classificadas"] += 1
                    register_report(idx, token, "fallback_root", "Alínea sem pai; criada como raiz")
                novo_item = ItemSessao.objects.create(
                    sessao=sessao_atual,
                    subsessao=subsessao_atual,
                    parent=parent,
                    enum_tipo=ItemSessao.EnumTipo.ALINEA if parent else ItemSessao.EnumTipo.NENHUM,
                    texto=data["texto"],
                    ordem=_next_ordem_item(sessao_atual, parent.id if parent else None),
                )
                counters["itens"] += 1
                counters["alineas"] += 1
                ultimo_item_valido = novo_item
                ultimo_item_estrutural = novo_item
                continue

            if token in {"OU_LINE", "BRACKET_NOTE"}:
                if sessao_atual is None:
                    ensure_sessao("0", idx)
                # Observações/alternativas ficam como filhos do último item estrutural válido.
                parent = ultimo_item_estrutural or ultimo_item_valido
                texto_comentario = _apply_import_comment_prefix(data["texto"])
                if parent is None and TEXT_IMPORT_STRICT_MODE:
                    counters["nao_classificadas"] += 1
                    register_report(idx, token, "fallback_root", "Comentário sem item anterior; criado como raiz")
                novo_item = ItemSessao.objects.create(
                    sessao=sessao_atual,
                    subsessao=subsessao_atual,
                    parent=parent,
                    enum_tipo=ItemSessao.EnumTipo.NENHUM,
                    texto=texto_comentario,
                    ordem=_next_ordem_item(sessao_atual, parent.id if parent else None),
                )
                counters["itens"] += 1
                counters["comentarios"] += 1
                ultimo_item_valido = novo_item
                continue

            if token == "PLAIN_TEXT":
                if sessao_atual is None:
                    sessao_atual = SessaoTermo.objects.create(
                        termo=termo,
                        titulo=_truncate_to_model_field("Conteudo importado", max_titulo_sessao),
                        ordem=_next_ordem_sessao(termo),
                    )
                    counters["sessoes"] += 1
                    sessao_numero_atual = None
                    if TEXT_IMPORT_STRICT_MODE:
                        register_report(idx, token, "create_default_section", "Texto sem sessão numerada prévia")
                if ultimo_item_valido is not None:
                    # Em texto corrido sem marcador, agrega ao item anterior para não perder conteúdo.
                    ultimo_item_valido.texto = f"{ultimo_item_valido.texto}\n{data['texto']}"
                    ultimo_item_valido.save(update_fields=["texto"])
                else:
                    if TEXT_IMPORT_STRICT_MODE:
                        counters["nao_classificadas"] += 1
                        register_report(idx, token, "fallback_root", "Texto sem item anterior; criado como item raiz")
                    novo_item = ItemSessao.objects.create(
                        sessao=sessao_atual,
                        subsessao=subsessao_atual,
                        parent=None,
                        enum_tipo=ItemSessao.EnumTipo.NENHUM,
                        texto=data["texto"],
                        ordem=_next_ordem_item(sessao_atual, None),
                    )
                    counters["itens"] += 1
                    ultimo_item_valido = novo_item
                    ultimo_item_estrutural = novo_item
                continue

        logger.info("Importação textual concluída para termo=%s resumo=%s", termo.pk, counters)
        if parse_report:
            for entry in parse_report:
                logger.info(
                    "Importação textual termo=%s linha=%s token=%s acao=%s detalhe=%s",
                    termo.pk,
                    entry["linha"],
                    entry["token"],
                    entry["acao"],
                    entry["detalhe"],
                )

        return termo, counters


def _importar_termo_docx(arquivo, apelido: str) -> TermoReferencia:
    """Função `_importar_termo_docx` no fluxo do app `licitacoes`.

    Objetivo pedagógico:
    - Executa uma etapa específica de validação, transformação ou orquestração.
    - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
    - Parâmetros: `arquivo`, `apelido`.
    """

    from docx import Document

    # Leitura em memória evita manipulação temporária no filesystem do servidor.
    payload = arquivo.read()
    documento = Document(BytesIO(payload))
    max_titulo_sessao = SessaoTermo._meta.get_field("titulo").max_length
    max_titulo_subsessao = SubsessaoTermo._meta.get_field("titulo").max_length

    with transaction.atomic():
        termo = TermoReferencia.objects.create(
            apelido=apelido,
            processo_sei="IMPORTADO DOCX",
            link_processo_sei="",
        )

        sessao_atual: SessaoTermo | None = None
        subsessao_atual: SubsessaoTermo | None = None
        item_nivel3_atual: ItemSessao | None = None

        # Parsing orientado por estilo e fallback por regex de enumeração textual.
        for paragraph in documento.paragraphs:
            texto_original = (paragraph.text or "").strip()
            if not texto_original:
                continue

            nivel = _docx_level_from_style(getattr(paragraph.style, "name", ""))
            texto_item, enum_tipo = _extract_item_text_and_enum(texto_original)
            if not texto_item:
                continue

            if nivel == 1:
                sessao_atual = SessaoTermo.objects.create(
                    termo=termo,
                    titulo=_truncate_to_model_field(texto_item, max_titulo_sessao),
                    ordem=_next_ordem_sessao(termo),
                )
                subsessao_atual = None
                item_nivel3_atual = None
                continue

            if sessao_atual is None:
                sessao_atual = SessaoTermo.objects.create(
                    termo=termo,
                    titulo=_truncate_to_model_field("Conteudo importado", max_titulo_sessao),
                    ordem=_next_ordem_sessao(termo),
                )

            if nivel == 2:
                subsessao_atual = SubsessaoTermo.objects.create(
                    sessao=sessao_atual,
                    titulo=_truncate_to_model_field(texto_item, max_titulo_subsessao),
                    ordem=_next_ordem_subsessao(sessao_atual),
                )
                item_nivel3_atual = None
                continue

            if nivel == 3:
                item_nivel3_atual = ItemSessao.objects.create(
                    sessao=sessao_atual,
                    subsessao=subsessao_atual,
                    parent=None,
                    enum_tipo=enum_tipo,
                    texto=texto_item,
                    ordem=_next_ordem_item(sessao_atual, None),
                )
                continue

            if nivel == 4 and _is_alternative_or_comment_line(texto_item):
                texto_item = _apply_import_comment_prefix(texto_item)

            if nivel is None and _is_alternative_or_comment_line(texto_item):
                texto_item = _apply_import_comment_prefix(texto_item)

            ItemSessao.objects.create(
                sessao=sessao_atual,
                subsessao=subsessao_atual,
                parent=item_nivel3_atual,
                enum_tipo=enum_tipo if item_nivel3_atual else ItemSessao.EnumTipo.NENHUM,
                texto=texto_item,
                ordem=_next_ordem_item(sessao_atual, item_nivel3_atual.id if item_nivel3_atual else None),
            )

        return termo


class LicitacoesAccessMixin(UserPassesTestMixin):
    """Classe `LicitacoesAccessMixin` do app `licitacoes`.

    Objetivo pedagógico:
    - Organiza comportamento e estado relacionados a uma parte do domínio.
    - Em auditoria, inspecione métodos públicos e pontos de escrita em banco.
    """

    def test_func(self) -> bool:
        """Método `LicitacoesAccessMixin.test_func` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`.
        """

        return _can_access_licitacoes(self.request.user)


class LicitacoesHomeView(LicitacoesAccessMixin, TemplateView):
    """Classe `LicitacoesHomeView` do app `licitacoes`.

    Objetivo pedagógico:
    - Organiza comportamento e estado relacionados a uma parte do domínio.
    - Em auditoria, inspecione métodos públicos e pontos de escrita em banco.
    """

    template_name = "licitacoes/home.html"


class EtpTicListView(LicitacoesAccessMixin, ListView):
    model = EtpTic
    template_name = "licitacoes/etp_tic_list.html"
    context_object_name = "etps"


class EtpTicCreateView(LicitacoesAccessMixin, CreateView):
    model = EtpTic
    form_class = EtpTicCreateForm
    template_name = "licitacoes/etp_tic_create.html"

    def form_valid(self, form):
        self.object = form.save(commit=False)
        self.object.criado_por = self.request.user
        self.object.atualizado_por = self.request.user
        self.object.declaracao_viabilidade = EtpTic.DECLARACAO_PADRAO
        self.object.save()
        return redirect("licitacoes_etp_tic_edit", pk=self.object.pk)


class EtpTicEditView(LicitacoesAccessMixin, UpdateView):
    model = EtpTic
    form_class = EtpTicSecaoForm
    template_name = "licitacoes/etp_tic_edit.html"
    context_object_name = "etp"

    def _get_secao_numero(self):
        secao = self.request.GET.get("secao")
        try:
            numero = int(secao) if secao is not None else int(self.object.secao_atual or 1)
        except (TypeError, ValueError):
            numero = 1
        if numero not in ETP_TIC_SECOES_MAP:
            numero = 1
        return numero

    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        numero = self._get_secao_numero()
        kwargs["section_fields"] = ETP_TIC_SECOES_MAP[numero]["campos"]
        return kwargs

    def form_valid(self, form):
        etp = form.save(commit=False)
        secao = self._get_secao_numero()
        acao = self.request.POST.get("_acao", "salvar")
        if acao == "proximo":
            etp.secao_atual = min(secao + 1, 18)
        elif acao == "anterior":
            etp.secao_atual = max(secao - 1, 1)
        else:
            etp.secao_atual = secao
        etp.atualizado_por = self.request.user
        etp.save()
        next_secao = etp.secao_atual
        return redirect(f"{reverse('licitacoes_etp_tic_edit', kwargs={'pk': etp.pk})}?secao={next_secao}")

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        secao_numero = self._get_secao_numero()
        status = _etp_status_por_secao(self.object)
        secoes_navegacao = []
        for secao in ETP_TIC_SECOES:
            secao_copy = dict(secao)
            secao_copy["status"] = status.get(secao["numero"], "nao_iniciado")
            secoes_navegacao.append(secao_copy)
        context["secao_atual"] = ETP_TIC_SECOES_MAP[secao_numero]
        context["secao_numero"] = secao_numero
        context["secoes"] = secoes_navegacao
        context["preview_url"] = reverse("licitacoes_etp_tic_preview", kwargs={"pk": self.object.pk})
        context["export_url"] = reverse("licitacoes_etp_tic_export_docx", kwargs={"pk": self.object.pk})
        return context


class EtpTicAutosaveView(LicitacoesAccessMixin, View):
    def post(self, request, pk: int):
        etp = get_object_or_404(EtpTic, pk=pk)
        try:
            payload = json.loads(request.body.decode("utf-8") if request.body else "{}")
        except Exception:
            return JsonResponse({"ok": False, "error": "payload invalido"}, status=400)
        try:
            secao = int(payload.get("secao") or 1)
        except (TypeError, ValueError):
            secao = 1
        if secao not in ETP_TIC_SECOES_MAP:
            return JsonResponse({"ok": False, "error": "secao invalida"}, status=400)
        campos = ETP_TIC_SECOES_MAP[secao]["campos"]
        data = {}
        for campo in campos:
            if campo == "declaracao_viabilidade":
                continue
            if campo in payload:
                data[campo] = payload[campo]
        form = EtpTicSecaoForm(data, instance=etp, section_fields=campos)
        if not form.is_valid():
            return JsonResponse({"ok": False, "errors": form.errors}, status=400)
        etp = form.save(commit=False)
        etp.secao_atual = secao
        etp.atualizado_por = request.user
        etp.save()
        status = _etp_status_por_secao(etp)
        concluidas = [k for k, v in status.items() if v == "concluido"]
        return JsonResponse(
            {
                "ok": True,
                "salvo_em": timezone.localtime(etp.atualizado_em).strftime("%d/%m/%Y %H:%M"),
                "progresso": round((len(concluidas) / 18) * 100, 2),
                "secoes_concluidas": concluidas,
            }
        )


class EtpTicPreviewView(LicitacoesAccessMixin, DetailView):
    model = EtpTic
    template_name = "licitacoes/etp_tic_preview.html"
    context_object_name = "etp"

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context["secoes_render"] = _etp_render_sections(self.object)
        return context


class EtpTicConcluirView(LicitacoesAccessMixin, View):
    def post(self, request, pk: int):
        etp = get_object_or_404(EtpTic, pk=pk)
        if not etp.numero_processo_servico.strip():
            messages.error(request, "Numero do processo servico e obrigatorio para concluir.")
            return redirect("licitacoes_etp_tic_edit", pk=pk)
        if not (etp.declaracao_viabilidade or "").strip():
            messages.error(request, "Declaracao de viabilidade e obrigatoria para concluir.")
            return redirect("licitacoes_etp_tic_edit", pk=pk)
        etp.status = EtpTic.Status.CONCLUIDO
        etp.atualizado_por = request.user
        etp.save(update_fields=["status", "atualizado_por", "atualizado_em"])
        messages.success(request, "ETP TIC concluido.")
        return redirect("licitacoes_etp_tic_preview", pk=pk)


class EtpTicDeleteView(LicitacoesAccessMixin, DeleteView):
    model = EtpTic
    template_name = "licitacoes/etp_tic_confirm_delete.html"
    success_url = reverse_lazy("licitacoes_etp_tic_list")


class EtpTicExportDocxView(LicitacoesAccessMixin, View):
    def get(self, request, pk: int):
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Cm, Pt
        except Exception:
            return HttpResponse(
                "Exportacao DOCX indisponivel. Instale a dependencia python-docx no servidor.",
                status=503,
                content_type="text/plain; charset=utf-8",
            )

        etp = get_object_or_404(EtpTic, pk=pk)
        secoes = _etp_render_sections(etp)

        document = Document()
        normal_style = document.styles["Normal"]
        normal_style.font.name = "Verdana"
        normal_style.font.size = Pt(10)

        section = document.sections[0]
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

        nome_etp = (etp.titulo or "").strip() or etp.numero_processo_servico
        p_titulo = document.add_paragraph()
        p_titulo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p_titulo.add_run(f"ETP TIC - {nome_etp}")
        run.bold = True
        run.font.name = "Verdana"
        run.font.size = Pt(11)
        comment_author = (
            request.user.get_full_name().strip()
            or request.user.username
            or "SGI SEDS"
        )
        comment_initials = "".join(
            parte[0].upper() for parte in comment_author.split()[:2] if parte
        ) or "SG"

        for secao in secoes:
            header = document.add_paragraph()
            header.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run_h = header.add_run(f"{secao['numero']}. {secao['titulo']}")
            run_h.bold = True
            run_h.font.name = "Verdana"
            run_h.font.size = Pt(10)

            if secao["descricao"]:
                # Comentário nativo do Word (balão lateral), ancorado no título da seção.
                document.add_comment(
                    run_h,
                    text=secao["descricao"],
                    author=comment_author,
                    initials=comment_initials,
                )

            if secao["entradas"]:
                for entrada in secao["entradas"]:
                    p = document.add_paragraph(entrada)
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            else:
                p_vazio = document.add_paragraph(f"{secao['numero']}.1. -")
                p_vazio.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            document.add_paragraph("")

        buffer = BytesIO()
        document.save(buffer)
        buffer.seek(0)
        filename = f"etp_tic_{slugify(etp.titulo or etp.numero_processo_servico) or etp.pk}.docx"
        response = HttpResponse(
            buffer.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        response["Content-Disposition"] = f'attachment; filename="{filename}"'
        return response


class TermoReferenciaListView(LicitacoesAccessMixin, ListView):
    """Classe `TermoReferenciaListView` do app `licitacoes`.

    Objetivo pedagógico:
    - Organiza comportamento e estado relacionados a uma parte do domínio.
    - Em auditoria, inspecione métodos públicos e pontos de escrita em banco.
    """

    model = TermoReferencia
    template_name = "licitacoes/termo_list.html"
    context_object_name = "termos"


class TermoReferenciaCreateView(LicitacoesAccessMixin, CreateView):
    """Classe `TermoReferenciaCreateView` do app `licitacoes`.

    Objetivo pedagógico:
    - Organiza comportamento e estado relacionados a uma parte do domínio.
    - Em auditoria, inspecione métodos públicos e pontos de escrita em banco.
    """

    model = TermoReferencia
    form_class = TermoReferenciaForm
    template_name = "licitacoes/termo_form.html"
    success_url = reverse_lazy("licitacoes_termo_list")


class TermoReferenciaImportarView(LicitacoesAccessMixin, FormView):
    """Classe `TermoReferenciaImportarView` do app `licitacoes`.

    Objetivo pedagógico:
    - Organiza comportamento e estado relacionados a uma parte do domínio.
    - Em auditoria, inspecione métodos públicos e pontos de escrita em banco.
    """

    template_name = "licitacoes/termo_importar_form.html"
    form_class = TermoReferenciaImportarForm

    def form_valid(self, form):
        """Método `TermoReferenciaImportarView.form_valid` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`, `form`.
        """

        apelido = form.cleaned_data["apelido"].strip()
        texto_colado = (form.cleaned_data.get("texto") or "").strip()
        arquivo = form.cleaned_data.get("arquivo")
        resumo: dict[str, int] = {}
        try:
            if texto_colado:
                termo, resumo = _importar_termo_texto(texto_colado, apelido)
            else:
                termo = _importar_termo_docx(arquivo=arquivo, apelido=apelido)
        except ModuleNotFoundError:
            form.add_error(
                "arquivo",
                "Importacao DOCX indisponivel. Instale a dependencia python-docx no servidor.",
            )
            return self.form_invalid(form)
        except Exception as exc:
            campo_erro = "texto" if texto_colado else "arquivo"
            form.add_error(campo_erro, f"Falha ao importar o documento: {exc}")
            return self.form_invalid(form)
        if resumo:
            messages.success(
                self.request,
                (
                    "Importacao concluida: "
                    f"{resumo.get('sessoes', 0)} secoes, "
                    f"{resumo.get('subsessoes', 0)} subsecoes, "
                    f"{resumo.get('itens', 0)} itens, "
                    f"{resumo.get('incisos', 0)} incisos, "
                    f"{resumo.get('alineas', 0)} alineas, "
                    f"{resumo.get('comentarios', 0)} comentarios."
                ),
            )
        return redirect("licitacoes_termo_detail", pk=termo.pk)


class TermoReferenciaUpdateView(LicitacoesAccessMixin, UpdateView):
    """Classe `TermoReferenciaUpdateView` do app `licitacoes`.

    Objetivo pedagógico:
    - Organiza comportamento e estado relacionados a uma parte do domínio.
    - Em auditoria, inspecione métodos públicos e pontos de escrita em banco.
    """

    model = TermoReferencia
    form_class = TermoReferenciaForm
    template_name = "licitacoes/termo_form.html"

    def get_success_url(self):
        """Método `TermoReferenciaUpdateView.get_success_url` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`.
        """

        return reverse("licitacoes_termo_detail", kwargs={"pk": self.object.pk})


class TermoReferenciaDeleteView(LicitacoesAccessMixin, DeleteView):
    """Classe `TermoReferenciaDeleteView` do app `licitacoes`.

    Objetivo pedagógico:
    - Organiza comportamento e estado relacionados a uma parte do domínio.
    - Em auditoria, inspecione métodos públicos e pontos de escrita em banco.
    """

    model = TermoReferencia
    template_name = "licitacoes/termo_confirm_delete.html"
    success_url = reverse_lazy("licitacoes_termo_list")


class TermoReferenciaDuplicarView(LicitacoesAccessMixin, FormView):
    """Classe `TermoReferenciaDuplicarView` do app `licitacoes`.

    Objetivo pedagógico:
    - Organiza comportamento e estado relacionados a uma parte do domínio.
    - Em auditoria, inspecione métodos públicos e pontos de escrita em banco.
    """

    template_name = "licitacoes/termo_duplicar_form.html"
    form_class = TermoReferenciaDuplicarForm

    def dispatch(self, request, *args, **kwargs):
        """Método `TermoReferenciaDuplicarView.dispatch` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`, `request`, `*args`, `**kwargs`.
        """

        self.origem = get_object_or_404(TermoReferencia, pk=kwargs["pk"])
        return super().dispatch(request, *args, **kwargs)

    def get_context_data(self, **kwargs):
        """Método `TermoReferenciaDuplicarView.get_context_data` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`, `**kwargs`.
        """

        context = super().get_context_data(**kwargs)
        context["termo_origem"] = self.origem
        return context

    def get_initial(self):
        """Método `TermoReferenciaDuplicarView.get_initial` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`.
        """

        initial = super().get_initial()
        initial["novo_apelido"] = f"{self.origem.apelido} (copia)"
        return initial

    def form_valid(self, form):
        """Método `TermoReferenciaDuplicarView.form_valid` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`, `form`.
        """

        novo_nome = form.cleaned_data["novo_apelido"].strip()
        novo_termo = _duplicar_estrutura_termo(self.origem, novo_nome)
        from django.shortcuts import redirect

        return redirect("licitacoes_termo_detail", pk=novo_termo.pk)


class TermoReferenciaDetailView(LicitacoesAccessMixin, DetailView):
    """Classe `TermoReferenciaDetailView` do app `licitacoes`.

    Objetivo pedagógico:
    - Organiza comportamento e estado relacionados a uma parte do domínio.
    - Em auditoria, inspecione métodos públicos e pontos de escrita em banco.
    """

    model = TermoReferencia
    template_name = "licitacoes/termo_detail.html"
    context_object_name = "termo"

    def get_context_data(self, **kwargs):
        """Método `TermoReferenciaDetailView.get_context_data` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`, `**kwargs`.
        """

        context = super().get_context_data(**kwargs)
        termo = context["termo"]
        # Pré-carrega todas as linhas de tabela do termo para renderizar sem N+1.
        tabela_linhas = (
            TabelaItemLinha.objects.filter(item__sessao__termo=termo)
            .select_related("item")
            .order_by("item_id", "ordem", "id")
        )
        tabela_por_item: dict[int, list[TabelaItemLinha]] = defaultdict(list)
        for linha in tabela_linhas:
            tabela_por_item[linha.item_id].append(linha)

        sessoes_data = []
        for sessao in termo.sessoes.order_by("ordem", "id"):
            itens_rows = _build_item_rows(sessao, sessao.ordem)
            subsessoes = list(sessao.subsessoes.order_by("ordem", "id"))
            primeiro_item_por_subsessao: dict[int, int] = {}
            # Descobre o primeiro item de cada subseção para inserir heading no ponto correto da lista.
            for row in itens_rows:
                if row["depth"] != 0:
                    continue
                sub_id = row["item"].subsessao_id
                if sub_id and sub_id not in primeiro_item_por_subsessao:
                    primeiro_item_por_subsessao[sub_id] = row["item"].id
            insercao_subsessao_por_item: dict[int, SubsessaoTermo] = {}
            for subsessao in subsessoes:
                first_item_id = primeiro_item_por_subsessao.get(subsessao.id)
                if first_item_id:
                    insercao_subsessao_por_item[first_item_id] = subsessao
            subsessoes_sem_itens = [
                subsessao
                for subsessao in subsessoes
                if subsessao.id not in primeiro_item_por_subsessao
            ]
            for row in itens_rows:
                row["pode_tabela_itens"] = row["indice"] == "1.1"
                row["tabela_linhas"] = tabela_por_item.get(row["item"].id, [])
                row["subsessao_heading"] = insercao_subsessao_por_item.get(row["item"].id)
            sessoes_data.append(
                {
                    "sessao": sessao,
                    "indice": sessao.ordem,
                    "itens": itens_rows,
                    "subsessoes": subsessoes,
                    "subsessoes_sem_itens": subsessoes_sem_itens,
                }
            )
        context["sessoes_data"] = sessoes_data
        return context


class TermoReferenciaExportDocxView(LicitacoesAccessMixin, View):
    """Classe `TermoReferenciaExportDocxView` do app `licitacoes`.

    Objetivo pedagógico:
    - Organiza comportamento e estado relacionados a uma parte do domínio.
    - Em auditoria, inspecione métodos públicos e pontos de escrita em banco.
    """

    def get(self, request, pk: int):
        """Método `TermoReferenciaExportDocxView.get` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`, `request`, `pk`.
        """

        # Dependência opcional: responde 503 com orientação quando python-docx não está disponível.
        try:
            from docx import Document
            from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn
            from docx.shared import Cm, Pt
        except Exception:
            return HttpResponse(
                "Exportacao DOCX indisponivel. Instale a dependencia python-docx no servidor.",
                status=503,
                content_type="text/plain; charset=utf-8",
            )

        termo = get_object_or_404(TermoReferencia, pk=pk)
        document = Document()
        normal_style = document.styles["Normal"]
        normal_style.font.name = "Verdana"
        normal_style.font.size = Pt(10)

        def _apply_heading_style(paragraph, level: int) -> None:
            """Função `TermoReferenciaExportDocxView.get._apply_heading_style` no fluxo do app `licitacoes`.

            Objetivo pedagógico:
            - Executa uma etapa específica de validação, transformação ou orquestração.
            - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
            - Parâmetros: `paragraph`, `level`.
            """

            candidates = [
                f"Heading {level}",
                f"Titulo {level}",
                f"Titulo {level}".replace("i", "í"),
                f"Título {level}",
            ]
            # Tenta estilos alternativos porque templates Word variam por idioma/instalação.
            for style_name in candidates:
                try:
                    paragraph.style = document.styles[style_name]
                    return
                except Exception:
                    continue

        section = document.sections[0]
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

        titulo = document.add_paragraph()
        run_titulo = titulo.add_run(
            f"Termo de Referencia - {_expand_seds_references(termo.apelido)}"
        )
        run_titulo.bold = True
        run_titulo.font.name = "Verdana"
        run_titulo.font.size = Pt(10)
        titulo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        meta = document.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        meta.add_run("Processo SEI: ").bold = True
        meta.add_run(_expand_seds_references(termo.processo_sei or "-"))
        if termo.link_processo_sei:
            meta_link = document.add_paragraph()
            meta_link.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            meta_link.add_run("Link do Processo SEI: ").bold = True
            meta_link.add_run(termo.link_processo_sei)

        document.add_paragraph("")

        for sessao in termo.sessoes.order_by("ordem", "id"):
            p_sessao = document.add_paragraph()
            p_sessao.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            _apply_heading_style(p_sessao, 1)
            r_sessao = p_sessao.add_run(
                f"{sessao.ordem}. {_expand_seds_references(sessao.titulo)}"
            )
            r_sessao.bold = True
            r_sessao.font.name = "Verdana"
            r_sessao.font.size = Pt(10)

            rows = _build_item_rows(sessao, sessao.ordem)
            subsessoes = list(sessao.subsessoes.order_by("ordem", "id"))
            primeiro_item_por_subsessao: dict[int, int] = {}
            for row in rows:
                if row["depth"] != 0:
                    continue
                sub_id = row["item"].subsessao_id
                if sub_id and sub_id not in primeiro_item_por_subsessao:
                    primeiro_item_por_subsessao[sub_id] = row["item"].id
            insercao_subsessao_por_item: dict[int, SubsessaoTermo] = {}
            for subsessao in subsessoes:
                first_item_id = primeiro_item_por_subsessao.get(subsessao.id)
                if first_item_id:
                    insercao_subsessao_por_item[first_item_id] = subsessao
            subsessoes_sem_itens = [
                subsessao
                for subsessao in subsessoes
                if subsessao.id not in primeiro_item_por_subsessao
            ]

            for row in rows:
                subsessao_heading = insercao_subsessao_por_item.get(row["item"].id)
                if subsessao_heading:
                    p_sub = document.add_paragraph()
                    p_sub.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p_sub.paragraph_format.left_indent = Cm(0.2)
                    _apply_heading_style(p_sub, 2)
                    r_sub = p_sub.add_run(_expand_seds_references(subsessao_heading.titulo))
                    r_sub.bold = True
                    r_sub.font.name = "Verdana"
                    r_sub.font.size = Pt(10)

                p_item = document.add_paragraph()
                p_item.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p_item.paragraph_format.left_indent = Cm(0.45 * row["depth"])
                if row["item"].enum_tipo in {ItemSessao.EnumTipo.INCISO, ItemSessao.EnumTipo.ALINEA}:
                    run_enum = p_item.add_run(f'{row["enum_prefix"]} ')
                    run_enum.bold = True
                    run_enum.font.name = "Verdana"
                    run_enum.font.size = Pt(10)
                else:
                    run_idx = p_item.add_run(f'{row["indice"]}. ')
                    run_idx.bold = True
                    run_idx.font.name = "Verdana"
                    run_idx.font.size = Pt(10)
                p_item.add_run(_expand_seds_references(row["item"].texto or ""))

                # A tabela de itens é materializada somente no item 1.1 por regra do módulo.
                if row["indice"] == "1.1":
                    linhas = list(row["item"].tabela_linhas.order_by("ordem", "id"))
                    if linhas:
                        doc_table = document.add_table(rows=1, cols=6)
                        doc_table.style = "Table Grid"
                        doc_table.alignment = WD_TABLE_ALIGNMENT.LEFT
                        doc_table.autofit = True
                        # Equivalente ao Word: Autoajuste > Ajustar-se automaticamente a janela.
                        tbl = doc_table._tbl
                        tbl_pr = tbl.tblPr
                        tbl_w = tbl_pr.find(qn("w:tblW"))
                        if tbl_w is None:
                            tbl_w = OxmlElement("w:tblW")
                            tbl_pr.append(tbl_w)
                        tbl_w.set(qn("w:type"), "pct")
                        tbl_w.set(qn("w:w"), "5000")

                        headers = [
                            "Item",
                            "Descricao",
                            "CATMAT/CATSER",
                            "Siafisico",
                            "UF",
                            "Quantidade",
                        ]
                        hdr_cells = doc_table.rows[0].cells
                        for idx, title in enumerate(headers):
                            hdr_p = hdr_cells[idx].paragraphs[0]
                            hdr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            hdr_p.paragraph_format.space_before = Pt(0)
                            hdr_p.paragraph_format.space_after = Pt(0)
                            hdr_p.paragraph_format.line_spacing = 1.0
                            hdr_run = hdr_p.add_run(title)
                            hdr_run.bold = True
                            hdr_run.font.name = "Verdana"
                            hdr_run.font.size = Pt(8)
                            hdr_cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

                        # Uniformiza altura das linhas com base na maior descricao.
                        # Aproximacao de quebra para Verdana 8 com autoajuste na janela.
                        def _line_count(text: str, chars_per_line: int = 38) -> int:
                            """Função `TermoReferenciaExportDocxView.get._line_count` no fluxo do app `licitacoes`.

                            Objetivo pedagógico:
                            - Executa uma etapa específica de validação, transformação ou orquestração.
                            - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
                            - Parâmetros: `text`, `chars_per_line`.
                            """

                            if not text:
                                return 1
                            parts = str(text).splitlines() or [str(text)]
                            total = 0
                            for part in parts:
                                total += max(1, (len(part) + chars_per_line - 1) // chars_per_line)
                            return max(total, 1)

                        max_lines = 1
                        for linha in linhas:
                            max_lines = max(max_lines, _line_count(linha.descricao))
                        # Altura unica para todas as linhas, baseada na maior.
                        row_height = Pt(12 + (max_lines * 11))
                        row_height_twips = str(int(round((12 + (max_lines * 11)) * 20)))

                        for linha in linhas:
                            row_cells = doc_table.add_row().cells
                            row_cells[0].text = str(linha.ordem)
                            row_cells[1].text = _expand_seds_references(linha.descricao or "-")
                            row_cells[2].text = linha.catmat_catser or "-"
                            row_cells[3].text = linha.siafisico or "-"
                            row_cells[4].text = linha.unidade_fornecimento or "-"
                            row_cells[5].text = _quantidade_text(linha.quantidade)
                            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                            row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            row_cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for cell in row_cells:
                                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                                for paragraph in cell.paragraphs:
                                    paragraph.paragraph_format.space_before = Pt(0)
                                    paragraph.paragraph_format.space_after = Pt(0)
                                    paragraph.paragraph_format.line_spacing = 1.0
                                    for run in paragraph.runs:
                                        run.font.name = "Verdana"
                                        run.font.size = Pt(8)

                        # Reaplica no final para garantir uniformidade em todas as linhas.
                        for table_row in doc_table.rows:
                            table_row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
                            table_row.height = row_height
                            tr_pr = table_row._tr.get_or_add_trPr()
                            tr_h = tr_pr.find(qn("w:trHeight"))
                            if tr_h is None:
                                tr_h = OxmlElement("w:trHeight")
                                tr_pr.append(tr_h)
                            tr_h.set(qn("w:val"), row_height_twips)
                            tr_h.set(qn("w:hRule"), "exact")
                        document.add_paragraph("")

            for subsessao in subsessoes_sem_itens:
                p_sub = document.add_paragraph()
                p_sub.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p_sub.paragraph_format.left_indent = Cm(0.2)
                _apply_heading_style(p_sub, 2)
                r_sub = p_sub.add_run(_expand_seds_references(subsessao.titulo))
                r_sub.bold = True
                r_sub.font.name = "Verdana"
                r_sub.font.size = Pt(10)

            document.add_paragraph("")

        output = BytesIO()
        document.save(output)
        output.seek(0)
        filename = f"termo_referencia_{slugify(termo.apelido) or termo.pk}.docx"
        response = HttpResponse(
            output.getvalue(),
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        response["Content-Disposition"] = f'attachment; filename="{filename}"'
        return response


class SessaoTermoCreateView(LicitacoesAccessMixin, CreateView):
    """Classe `SessaoTermoCreateView` do app `licitacoes`.

    Objetivo pedagógico:
    - Organiza comportamento e estado relacionados a uma parte do domínio.
    - Em auditoria, inspecione métodos públicos e pontos de escrita em banco.
    """

    model = SessaoTermo
    form_class = SessaoTermoForm
    template_name = "licitacoes/sessao_form.html"

    def dispatch(self, request, *args, **kwargs):
        """Método `SessaoTermoCreateView.dispatch` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`, `request`, `*args`, `**kwargs`.
        """

        self.termo = get_object_or_404(TermoReferencia, pk=kwargs["termo_pk"])
        return super().dispatch(request, *args, **kwargs)

    def get_context_data(self, **kwargs):
        """Método `SessaoTermoCreateView.get_context_data` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`, `**kwargs`.
        """

        context = super().get_context_data(**kwargs)
        context["termo"] = self.termo
        context["next_url"] = _safe_next_url(
            self.request,
            reverse("licitacoes_termo_detail", kwargs={"pk": self.termo.pk}),
        )
        return context

    def form_valid(self, form):
        """Método `SessaoTermoCreateView.form_valid` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`, `form`.
        """

        form.instance.termo = self.termo
        form.instance.ordem = _next_ordem_sessao(self.termo)
        return super().form_valid(form)

    def get_success_url(self):
        """Método `SessaoTermoCreateView.get_success_url` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`.
        """

        default_url = reverse("licitacoes_termo_detail", kwargs={"pk": self.termo.pk})
        return _safe_next_url(self.request, default_url)


class SessaoTermoUpdateView(LicitacoesAccessMixin, UpdateView):
    """Classe `SessaoTermoUpdateView` do app `licitacoes`.

    Objetivo pedagógico:
    - Organiza comportamento e estado relacionados a uma parte do domínio.
    - Em auditoria, inspecione métodos públicos e pontos de escrita em banco.
    """

    model = SessaoTermo
    form_class = SessaoTermoForm
    template_name = "licitacoes/sessao_form.html"

    def dispatch(self, request, *args, **kwargs):
        """Método `SessaoTermoUpdateView.dispatch` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`, `request`, `*args`, `**kwargs`.
        """

        self.termo = get_object_or_404(TermoReferencia, pk=kwargs["termo_pk"])
        return super().dispatch(request, *args, **kwargs)

    def get_queryset(self):
        """Método `SessaoTermoUpdateView.get_queryset` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`.
        """

        return SessaoTermo.objects.filter(termo=self.termo)

    def get_context_data(self, **kwargs):
        """Método `SessaoTermoUpdateView.get_context_data` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`, `**kwargs`.
        """

        context = super().get_context_data(**kwargs)
        context["termo"] = self.termo
        context["next_url"] = _safe_next_url(
            self.request,
            reverse("licitacoes_termo_detail", kwargs={"pk": self.termo.pk}),
        )
        return context

    def get_success_url(self):
        """Método `SessaoTermoUpdateView.get_success_url` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`.
        """

        default_url = reverse("licitacoes_termo_detail", kwargs={"pk": self.termo.pk})
        return _safe_next_url(self.request, default_url)


class SessaoTermoDeleteView(LicitacoesAccessMixin, DeleteView):
    """Classe `SessaoTermoDeleteView` do app `licitacoes`.

    Objetivo pedagógico:
    - Organiza comportamento e estado relacionados a uma parte do domínio.
    - Em auditoria, inspecione métodos públicos e pontos de escrita em banco.
    """

    model = SessaoTermo
    template_name = "licitacoes/sessao_confirm_delete.html"

    def dispatch(self, request, *args, **kwargs):
        """Método `SessaoTermoDeleteView.dispatch` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`, `request`, `*args`, `**kwargs`.
        """

        self.termo = get_object_or_404(TermoReferencia, pk=kwargs["termo_pk"])
        return super().dispatch(request, *args, **kwargs)

    def get_queryset(self):
        """Método `SessaoTermoDeleteView.get_queryset` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`.
        """

        return SessaoTermo.objects.filter(termo=self.termo)

    def get_context_data(self, **kwargs):
        """Método `SessaoTermoDeleteView.get_context_data` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`, `**kwargs`.
        """

        context = super().get_context_data(**kwargs)
        context["next_url"] = _safe_next_url(
            self.request,
            reverse("licitacoes_termo_detail", kwargs={"pk": self.termo.pk}),
        )
        return context

    def form_valid(self, form):
        """Método `SessaoTermoDeleteView.form_valid` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`, `form`.
        """

        response = super().form_valid(form)
        _normalize_sessoes(self.termo)
        return response

    def get_success_url(self):
        """Método `SessaoTermoDeleteView.get_success_url` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`.
        """

        default_url = reverse("licitacoes_termo_detail", kwargs={"pk": self.termo.pk})
        return _safe_next_url(self.request, default_url)


class SessaoMoveView(LicitacoesAccessMixin, View):
    """Classe `SessaoMoveView` do app `licitacoes`.

    Objetivo pedagógico:
    - Organiza comportamento e estado relacionados a uma parte do domínio.
    - Em auditoria, inspecione métodos públicos e pontos de escrita em banco.
    """

    direction = "up"

    def post(self, request, termo_pk: int, pk: int):
        """Método `SessaoMoveView.post` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`, `request`, `termo_pk`, `pk`.
        """

        termo = get_object_or_404(TermoReferencia, pk=termo_pk)
        sessao = get_object_or_404(SessaoTermo, pk=pk, termo=termo)
        with transaction.atomic():
            if self.direction == "up":
                alvo = (
                    termo.sessoes.filter(ordem__lt=sessao.ordem)
                    .order_by("-ordem", "-id")
                    .first()
                )
            else:
                alvo = (
                    termo.sessoes.filter(ordem__gt=sessao.ordem)
                    .order_by("ordem", "id")
                    .first()
                )
            if alvo:
                sessao.ordem, alvo.ordem = alvo.ordem, sessao.ordem
                sessao.save(update_fields=["ordem"])
                alvo.save(update_fields=["ordem"])
            _normalize_sessoes(termo)
        return self._redirect(request, termo.pk)

    def _redirect(self, request, termo_pk: int):
        """Método `SessaoMoveView._redirect` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`, `request`, `termo_pk`.
        """

        from django.shortcuts import redirect

        default_url = reverse("licitacoes_termo_detail", kwargs={"pk": termo_pk})
        return redirect(_safe_next_url(request, default_url))


class SessaoMoveUpView(SessaoMoveView):
    """Classe `SessaoMoveUpView` do app `licitacoes`.

    Objetivo pedagógico:
    - Organiza comportamento e estado relacionados a uma parte do domínio.
    - Em auditoria, inspecione métodos públicos e pontos de escrita em banco.
    """

    direction = "up"


class SessaoMoveDownView(SessaoMoveView):
    """Classe `SessaoMoveDownView` do app `licitacoes`.

    Objetivo pedagógico:
    - Organiza comportamento e estado relacionados a uma parte do domínio.
    - Em auditoria, inspecione métodos públicos e pontos de escrita em banco.
    """

    direction = "down"


class SubsessaoTermoCreateView(LicitacoesAccessMixin, CreateView):
    """Classe `SubsessaoTermoCreateView` do app `licitacoes`.

    Objetivo pedagógico:
    - Organiza comportamento e estado relacionados a uma parte do domínio.
    - Em auditoria, inspecione métodos públicos e pontos de escrita em banco.
    """

    model = SubsessaoTermo
    form_class = SubsessaoTermoForm
    template_name = "licitacoes/subsessao_form.html"

    def dispatch(self, request, *args, **kwargs):
        """Método `SubsessaoTermoCreateView.dispatch` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`, `request`, `*args`, `**kwargs`.
        """

        self.sessao = get_object_or_404(SessaoTermo, pk=kwargs["sessao_pk"])
        return super().dispatch(request, *args, **kwargs)

    def get_context_data(self, **kwargs):
        """Método `SubsessaoTermoCreateView.get_context_data` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`, `**kwargs`.
        """

        context = super().get_context_data(**kwargs)
        context["sessao"] = self.sessao
        context["termo"] = self.sessao.termo
        context["next_url"] = _safe_next_url(
            self.request,
            reverse("licitacoes_termo_detail", kwargs={"pk": self.sessao.termo.pk}),
        )
        return context

    def form_valid(self, form):
        """Método `SubsessaoTermoCreateView.form_valid` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`, `form`.
        """

        form.instance.sessao = self.sessao
        form.instance.ordem = _next_ordem_subsessao(self.sessao)
        return super().form_valid(form)

    def get_success_url(self):
        """Método `SubsessaoTermoCreateView.get_success_url` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`.
        """

        default_url = reverse("licitacoes_termo_detail", kwargs={"pk": self.sessao.termo.pk})
        return _safe_next_url(self.request, default_url)


class SubsessaoTermoUpdateView(LicitacoesAccessMixin, UpdateView):
    """Classe `SubsessaoTermoUpdateView` do app `licitacoes`.

    Objetivo pedagógico:
    - Organiza comportamento e estado relacionados a uma parte do domínio.
    - Em auditoria, inspecione métodos públicos e pontos de escrita em banco.
    """

    model = SubsessaoTermo
    form_class = SubsessaoTermoForm
    template_name = "licitacoes/subsessao_form.html"

    def dispatch(self, request, *args, **kwargs):
        """Método `SubsessaoTermoUpdateView.dispatch` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`, `request`, `*args`, `**kwargs`.
        """

        self.sessao = get_object_or_404(SessaoTermo, pk=kwargs["sessao_pk"])
        return super().dispatch(request, *args, **kwargs)

    def get_queryset(self):
        """Método `SubsessaoTermoUpdateView.get_queryset` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`.
        """

        return SubsessaoTermo.objects.filter(sessao=self.sessao)

    def get_context_data(self, **kwargs):
        """Método `SubsessaoTermoUpdateView.get_context_data` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`, `**kwargs`.
        """

        context = super().get_context_data(**kwargs)
        context["sessao"] = self.sessao
        context["termo"] = self.sessao.termo
        context["next_url"] = _safe_next_url(
            self.request,
            reverse("licitacoes_termo_detail", kwargs={"pk": self.sessao.termo.pk}),
        )
        return context

    def get_success_url(self):
        """Método `SubsessaoTermoUpdateView.get_success_url` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`.
        """

        default_url = reverse("licitacoes_termo_detail", kwargs={"pk": self.sessao.termo.pk})
        return _safe_next_url(self.request, default_url)


class SubsessaoTermoDeleteView(LicitacoesAccessMixin, DeleteView):
    """Classe `SubsessaoTermoDeleteView` do app `licitacoes`.

    Objetivo pedagógico:
    - Organiza comportamento e estado relacionados a uma parte do domínio.
    - Em auditoria, inspecione métodos públicos e pontos de escrita em banco.
    """

    model = SubsessaoTermo
    template_name = "licitacoes/subsessao_confirm_delete.html"

    def dispatch(self, request, *args, **kwargs):
        """Método `SubsessaoTermoDeleteView.dispatch` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`, `request`, `*args`, `**kwargs`.
        """

        self.sessao = get_object_or_404(SessaoTermo, pk=kwargs["sessao_pk"])
        return super().dispatch(request, *args, **kwargs)

    def get_queryset(self):
        """Método `SubsessaoTermoDeleteView.get_queryset` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`.
        """

        return SubsessaoTermo.objects.filter(sessao=self.sessao)

    def get_context_data(self, **kwargs):
        """Método `SubsessaoTermoDeleteView.get_context_data` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`, `**kwargs`.
        """

        context = super().get_context_data(**kwargs)
        context["next_url"] = _safe_next_url(
            self.request,
            reverse("licitacoes_termo_detail", kwargs={"pk": self.sessao.termo.pk}),
        )
        return context

    def form_valid(self, form):
        """Método `SubsessaoTermoDeleteView.form_valid` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`, `form`.
        """

        response = super().form_valid(form)
        _normalize_subsessoes(self.sessao)
        return response

    def get_success_url(self):
        """Método `SubsessaoTermoDeleteView.get_success_url` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`.
        """

        default_url = reverse("licitacoes_termo_detail", kwargs={"pk": self.sessao.termo.pk})
        return _safe_next_url(self.request, default_url)


class SubsessaoMoveView(LicitacoesAccessMixin, View):
    """Classe `SubsessaoMoveView` do app `licitacoes`.

    Objetivo pedagógico:
    - Organiza comportamento e estado relacionados a uma parte do domínio.
    - Em auditoria, inspecione métodos públicos e pontos de escrita em banco.
    """

    direction = "up"

    def post(self, request, sessao_pk: int, pk: int):
        """Método `SubsessaoMoveView.post` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`, `request`, `sessao_pk`, `pk`.
        """

        sessao = get_object_or_404(SessaoTermo, pk=sessao_pk)
        subsessao = get_object_or_404(SubsessaoTermo, pk=pk, sessao=sessao)
        with transaction.atomic():
            if self.direction == "up":
                alvo = (
                    sessao.subsessoes.filter(ordem__lt=subsessao.ordem)
                    .order_by("-ordem", "-id")
                    .first()
                )
            else:
                alvo = (
                    sessao.subsessoes.filter(ordem__gt=subsessao.ordem)
                    .order_by("ordem", "id")
                    .first()
                )
            if alvo:
                subsessao.ordem, alvo.ordem = alvo.ordem, subsessao.ordem
                subsessao.save(update_fields=["ordem"])
                alvo.save(update_fields=["ordem"])
            _normalize_subsessoes(sessao)
        from django.shortcuts import redirect

        default_url = reverse("licitacoes_termo_detail", kwargs={"pk": sessao.termo.pk})
        return redirect(_safe_next_url(request, default_url))


class SubsessaoMoveUpView(SubsessaoMoveView):
    """Classe `SubsessaoMoveUpView` do app `licitacoes`.

    Objetivo pedagógico:
    - Organiza comportamento e estado relacionados a uma parte do domínio.
    - Em auditoria, inspecione métodos públicos e pontos de escrita em banco.
    """

    direction = "up"


class SubsessaoMoveDownView(SubsessaoMoveView):
    """Classe `SubsessaoMoveDownView` do app `licitacoes`.

    Objetivo pedagógico:
    - Organiza comportamento e estado relacionados a uma parte do domínio.
    - Em auditoria, inspecione métodos públicos e pontos de escrita em banco.
    """

    direction = "down"


class ItemSessaoCreateView(LicitacoesAccessMixin, CreateView):
    """Classe `ItemSessaoCreateView` do app `licitacoes`.

    Objetivo pedagógico:
    - Organiza comportamento e estado relacionados a uma parte do domínio.
    - Em auditoria, inspecione métodos públicos e pontos de escrita em banco.
    """

    model = ItemSessao
    form_class = ItemSessaoForm
    template_name = "licitacoes/item_form.html"

    def dispatch(self, request, *args, **kwargs):
        """Método `ItemSessaoCreateView.dispatch` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`, `request`, `*args`, `**kwargs`.
        """

        self.sessao = get_object_or_404(SessaoTermo, pk=kwargs["sessao_pk"])
        self.parent = None
        self.subsessao = None
        self.enum_kind = request.GET.get("enum")
        subsessao_pk = kwargs.get("subsessao_pk")
        if subsessao_pk is not None:
            self.subsessao = get_object_or_404(SubsessaoTermo, pk=subsessao_pk, sessao=self.sessao)
        parent_pk = kwargs.get("parent_pk")
        if parent_pk is not None:
            self.parent = get_object_or_404(ItemSessao, pk=parent_pk, sessao=self.sessao)
            self.subsessao = self.parent.subsessao
        return super().dispatch(request, *args, **kwargs)

    def get_initial(self):
        """Método `ItemSessaoCreateView.get_initial` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`.
        """

        return super().get_initial()

    def get_form_kwargs(self):
        """Método `ItemSessaoCreateView.get_form_kwargs` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`.
        """

        kwargs = super().get_form_kwargs()
        kwargs["sessao"] = self.sessao
        return kwargs

    def get_context_data(self, **kwargs):
        """Método `ItemSessaoCreateView.get_context_data` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`, `**kwargs`.
        """

        context = super().get_context_data(**kwargs)
        context["sessao"] = self.sessao
        context["termo"] = self.sessao.termo
        context["parent_item"] = self.parent
        context["subsessao"] = self.subsessao
        context["enum_kind"] = self.enum_kind
        context["next_url"] = _safe_next_url(
            self.request,
            reverse("licitacoes_termo_detail", kwargs={"pk": self.sessao.termo.pk}),
        )
        return context

    def form_valid(self, form):
        """Método `ItemSessaoCreateView.form_valid` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`, `form`.
        """

        raw_text = (self.request.POST.get("texto") or "").strip()
        form.instance.sessao = self.sessao
        form.instance.parent = self.parent
        if self.parent is not None:
            form.instance.subsessao = self.parent.subsessao
        elif self.subsessao is not None:
            form.instance.subsessao = self.subsessao
        if re.match(r"^\d+(?:\.\d+)+(?:\.)?\s+", raw_text):
            form.instance.enum_tipo = ItemSessao.EnumTipo.NENHUM
        elif self.enum_kind == "inciso":
            form.instance.enum_tipo = ItemSessao.EnumTipo.INCISO
        elif self.enum_kind == "alinea":
            form.instance.enum_tipo = ItemSessao.EnumTipo.ALINEA
        else:
            form.instance.enum_tipo = ItemSessao.EnumTipo.NENHUM
        form.instance.ordem = _next_ordem_item(self.sessao, self.parent.id if self.parent else None)
        return super().form_valid(form)

    def get_success_url(self):
        """Método `ItemSessaoCreateView.get_success_url` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`.
        """

        default_url = reverse("licitacoes_termo_detail", kwargs={"pk": self.sessao.termo.pk})
        return _safe_next_url(self.request, default_url)


class ItemSessaoUpdateView(LicitacoesAccessMixin, UpdateView):
    """Classe `ItemSessaoUpdateView` do app `licitacoes`.

    Objetivo pedagógico:
    - Organiza comportamento e estado relacionados a uma parte do domínio.
    - Em auditoria, inspecione métodos públicos e pontos de escrita em banco.
    """

    model = ItemSessao
    form_class = ItemSessaoForm
    template_name = "licitacoes/item_form.html"

    def dispatch(self, request, *args, **kwargs):
        """Método `ItemSessaoUpdateView.dispatch` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`, `request`, `*args`, `**kwargs`.
        """

        self.sessao = get_object_or_404(SessaoTermo, pk=kwargs["sessao_pk"])
        return super().dispatch(request, *args, **kwargs)

    def get_form_kwargs(self):
        """Método `ItemSessaoUpdateView.get_form_kwargs` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`.
        """

        kwargs = super().get_form_kwargs()
        kwargs["sessao"] = self.sessao
        return kwargs

    def get_queryset(self):
        """Método `ItemSessaoUpdateView.get_queryset` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`.
        """

        return ItemSessao.objects.filter(sessao=self.sessao)

    def get_context_data(self, **kwargs):
        """Método `ItemSessaoUpdateView.get_context_data` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`, `**kwargs`.
        """

        context = super().get_context_data(**kwargs)
        context["sessao"] = self.sessao
        context["termo"] = self.sessao.termo
        context["parent_item"] = self.object.parent
        context["subsessao"] = self.object.subsessao
        context["next_url"] = _safe_next_url(
            self.request,
            reverse("licitacoes_termo_detail", kwargs={"pk": self.sessao.termo.pk}),
        )
        return context

    def form_valid(self, form):
        """Método `ItemSessaoUpdateView.form_valid` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`, `form`.
        """

        raw_text = (self.request.POST.get("texto") or "").strip()
        if self.object.parent_id:
            form.instance.subsessao = self.object.parent.subsessao
            if re.match(r"^\d+(?:\.\d+)+(?:\.)?\s+", raw_text):
                form.instance.enum_tipo = ItemSessao.EnumTipo.NENHUM
            else:
                form.instance.enum_tipo = self.object.enum_tipo
        return super().form_valid(form)

    def get_success_url(self):
        """Método `ItemSessaoUpdateView.get_success_url` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`.
        """

        default_url = reverse("licitacoes_termo_detail", kwargs={"pk": self.sessao.termo.pk})
        return _safe_next_url(self.request, default_url)


class ItemSessaoDeleteView(LicitacoesAccessMixin, DeleteView):
    """Classe `ItemSessaoDeleteView` do app `licitacoes`.

    Objetivo pedagógico:
    - Organiza comportamento e estado relacionados a uma parte do domínio.
    - Em auditoria, inspecione métodos públicos e pontos de escrita em banco.
    """

    model = ItemSessao
    template_name = "licitacoes/item_confirm_delete.html"

    def dispatch(self, request, *args, **kwargs):
        """Método `ItemSessaoDeleteView.dispatch` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`, `request`, `*args`, `**kwargs`.
        """

        self.sessao = get_object_or_404(SessaoTermo, pk=kwargs["sessao_pk"])
        return super().dispatch(request, *args, **kwargs)

    def get_queryset(self):
        """Método `ItemSessaoDeleteView.get_queryset` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`.
        """

        return ItemSessao.objects.filter(sessao=self.sessao)

    def get_context_data(self, **kwargs):
        """Método `ItemSessaoDeleteView.get_context_data` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`, `**kwargs`.
        """

        context = super().get_context_data(**kwargs)
        context["next_url"] = _safe_next_url(
            self.request,
            reverse("licitacoes_termo_detail", kwargs={"pk": self.sessao.termo.pk}),
        )
        return context

    def post(self, request, *args, **kwargs):
        """Torna a exclusão idempotente quando o mesmo POST é reenviado."""

        try:
            return super().post(request, *args, **kwargs)
        except Http404:
            return redirect(self.get_success_url())

    def form_valid(self, form):
        """Método `ItemSessaoDeleteView.form_valid` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`, `form`.
        """

        parent_id = self.object.parent_id
        response = super().form_valid(form)
        _normalize_items(self.sessao, parent_id)
        return response

    def get_success_url(self):
        """Método `ItemSessaoDeleteView.get_success_url` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`.
        """

        default_url = reverse("licitacoes_termo_detail", kwargs={"pk": self.sessao.termo.pk})
        return _safe_next_url(self.request, default_url)


class ItemMoveView(LicitacoesAccessMixin, View):
    """Classe `ItemMoveView` do app `licitacoes`.

    Objetivo pedagógico:
    - Organiza comportamento e estado relacionados a uma parte do domínio.
    - Em auditoria, inspecione métodos públicos e pontos de escrita em banco.
    """

    direction = "up"

    def post(self, request, sessao_pk: int, pk: int):
        """Método `ItemMoveView.post` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`, `request`, `sessao_pk`, `pk`.
        """

        sessao = get_object_or_404(SessaoTermo, pk=sessao_pk)
        item = get_object_or_404(ItemSessao, pk=pk, sessao=sessao)
        with transaction.atomic():
            if self.direction == "up":
                self._move_up(sessao, item)
            else:
                self._move_down(sessao, item)
        from django.shortcuts import redirect

        default_url = reverse("licitacoes_termo_detail", kwargs={"pk": sessao.termo.pk})
        return redirect(_safe_next_url(request, default_url))

    def _move_up(self, sessao: SessaoTermo, item: ItemSessao) -> None:
        """Método `ItemMoveView._move_up` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`, `sessao`, `item`.
        """

        parent_id = item.parent_id
        siblings = _ordered_item_ids(sessao, parent_id)
        pos = siblings.index(item.id)

        if parent_id is None and pos > 0:
            # Na raiz, seta para cima pode "entrar" no nivel do item anterior.
            prev_root = ItemSessao.objects.get(pk=siblings[pos - 1])
            old_group = [item_id for item_id in siblings if item_id != item.id]
            child_group = _ordered_item_ids(sessao, prev_root.id)
            item.parent = prev_root
            item.subsessao = prev_root.subsessao
            item.save(update_fields=["parent", "subsessao"])
            _apply_group_order(old_group)
            _apply_group_order(child_group + [item.id])
            return

        if pos > 0:
            siblings[pos - 1], siblings[pos] = siblings[pos], siblings[pos - 1]
            _apply_group_order(siblings)
            return

        if item.parent_id is None:
            return

        # Primeiro item do grupo: sobe um nivel (vira irmao do pai, antes dele).
        pai = item.parent
        assert pai is not None
        novo_parent_id = pai.parent_id
        grupo_origem = _ordered_item_ids(sessao, pai.id, exclude_id=item.id)
        grupo_destino = _ordered_item_ids(sessao, novo_parent_id)
        idx_pai = grupo_destino.index(pai.id)
        grupo_destino.insert(idx_pai, item.id)
        item.parent_id = novo_parent_id
        item.subsessao = pai.subsessao
        item.save(update_fields=["parent", "subsessao"])
        _apply_group_order(grupo_origem)
        _apply_group_order(grupo_destino)

    def _move_down(self, sessao: SessaoTermo, item: ItemSessao) -> None:
        """Método `ItemMoveView._move_down` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`, `sessao`, `item`.
        """

        parent_id = item.parent_id
        siblings = _ordered_item_ids(sessao, parent_id)
        pos = siblings.index(item.id)

        if parent_id is None and pos < len(siblings) - 1:
            # Na raiz, seta para baixo pode "entrar" no nivel do proximo item.
            next_root = ItemSessao.objects.get(pk=siblings[pos + 1])
            old_group = [item_id for item_id in siblings if item_id != item.id]
            child_group = _ordered_item_ids(sessao, next_root.id)
            item.parent = next_root
            item.subsessao = next_root.subsessao
            item.save(update_fields=["parent", "subsessao"])
            _apply_group_order(old_group)
            _apply_group_order([item.id] + child_group)
            return

        if pos < len(siblings) - 1:
            siblings[pos], siblings[pos + 1] = siblings[pos + 1], siblings[pos]
            _apply_group_order(siblings)
            return

        if item.parent_id is None:
            return

        # Ultimo item do grupo: desce um nivel (vira irmao do pai, depois dele).
        pai = item.parent
        assert pai is not None
        novo_parent_id = pai.parent_id
        grupo_origem = _ordered_item_ids(sessao, pai.id, exclude_id=item.id)
        grupo_destino = _ordered_item_ids(sessao, novo_parent_id)
        idx_pai = grupo_destino.index(pai.id)
        grupo_destino.insert(idx_pai + 1, item.id)
        item.parent_id = novo_parent_id
        item.subsessao = pai.subsessao
        item.save(update_fields=["parent", "subsessao"])
        _apply_group_order(grupo_origem)
        _apply_group_order(grupo_destino)


class ItemMoveUpView(ItemMoveView):
    """Classe `ItemMoveUpView` do app `licitacoes`.

    Objetivo pedagógico:
    - Organiza comportamento e estado relacionados a uma parte do domínio.
    - Em auditoria, inspecione métodos públicos e pontos de escrita em banco.
    """

    direction = "up"


class ItemMoveDownView(ItemMoveView):
    """Classe `ItemMoveDownView` do app `licitacoes`.

    Objetivo pedagógico:
    - Organiza comportamento e estado relacionados a uma parte do domínio.
    - Em auditoria, inspecione métodos públicos e pontos de escrita em banco.
    """

    direction = "down"


class TabelaItemLinhaCreateView(LicitacoesAccessMixin, CreateView):
    """Classe `TabelaItemLinhaCreateView` do app `licitacoes`.

    Objetivo pedagógico:
    - Organiza comportamento e estado relacionados a uma parte do domínio.
    - Em auditoria, inspecione métodos públicos e pontos de escrita em banco.
    """

    model = TabelaItemLinha
    form_class = TabelaItemLinhaForm
    template_name = "licitacoes/tabela_item_form.html"

    def dispatch(self, request, *args, **kwargs):
        """Método `TabelaItemLinhaCreateView.dispatch` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`, `request`, `*args`, `**kwargs`.
        """

        self.sessao = get_object_or_404(SessaoTermo, pk=kwargs["sessao_pk"])
        self.item = get_object_or_404(ItemSessao, pk=kwargs["item_pk"], sessao=self.sessao)
        if not _can_have_tabela_itens(self.item):
            raise Http404("Tabela de itens permitida somente para o item 1.1.")
        return super().dispatch(request, *args, **kwargs)

    def get_context_data(self, **kwargs):
        """Método `TabelaItemLinhaCreateView.get_context_data` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`, `**kwargs`.
        """

        context = super().get_context_data(**kwargs)
        context["sessao"] = self.sessao
        context["item"] = self.item
        context["termo"] = self.sessao.termo
        context["item_indice"] = _item_indice(self.item)
        context["next_url"] = _safe_next_url(
            self.request,
            f'{reverse("licitacoes_termo_detail", kwargs={"pk": self.sessao.termo.pk})}#tabela-itens-{self.item.pk}',
        )
        return context

    def form_valid(self, form):
        """Método `TabelaItemLinhaCreateView.form_valid` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`, `form`.
        """

        form.instance.item = self.item
        form.instance.ordem = _next_ordem_tabela(self.item)
        return super().form_valid(form)

    def get_success_url(self):
        """Método `TabelaItemLinhaCreateView.get_success_url` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`.
        """

        default_url = f'{reverse("licitacoes_termo_detail", kwargs={"pk": self.sessao.termo.pk})}#tabela-itens-{self.item.pk}'
        return _safe_next_url(self.request, default_url)


class TabelaItemLinhaUpdateView(LicitacoesAccessMixin, UpdateView):
    """Classe `TabelaItemLinhaUpdateView` do app `licitacoes`.

    Objetivo pedagógico:
    - Organiza comportamento e estado relacionados a uma parte do domínio.
    - Em auditoria, inspecione métodos públicos e pontos de escrita em banco.
    """

    model = TabelaItemLinha
    form_class = TabelaItemLinhaForm
    template_name = "licitacoes/tabela_item_form.html"

    def dispatch(self, request, *args, **kwargs):
        """Método `TabelaItemLinhaUpdateView.dispatch` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`, `request`, `*args`, `**kwargs`.
        """

        self.sessao = get_object_or_404(SessaoTermo, pk=kwargs["sessao_pk"])
        self.item = get_object_or_404(ItemSessao, pk=kwargs["item_pk"], sessao=self.sessao)
        return super().dispatch(request, *args, **kwargs)

    def get_queryset(self):
        """Método `TabelaItemLinhaUpdateView.get_queryset` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`.
        """

        return TabelaItemLinha.objects.filter(item=self.item)

    def get_context_data(self, **kwargs):
        """Método `TabelaItemLinhaUpdateView.get_context_data` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`, `**kwargs`.
        """

        context = super().get_context_data(**kwargs)
        context["sessao"] = self.sessao
        context["item"] = self.item
        context["termo"] = self.sessao.termo
        context["item_indice"] = _item_indice(self.item)
        context["next_url"] = _safe_next_url(
            self.request,
            f'{reverse("licitacoes_termo_detail", kwargs={"pk": self.sessao.termo.pk})}#tabela-itens-{self.item.pk}',
        )
        return context

    def get_success_url(self):
        """Método `TabelaItemLinhaUpdateView.get_success_url` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`.
        """

        default_url = f'{reverse("licitacoes_termo_detail", kwargs={"pk": self.sessao.termo.pk})}#tabela-itens-{self.item.pk}'
        return _safe_next_url(self.request, default_url)


class TabelaItemLinhaDeleteView(LicitacoesAccessMixin, DeleteView):
    """Classe `TabelaItemLinhaDeleteView` do app `licitacoes`.

    Objetivo pedagógico:
    - Organiza comportamento e estado relacionados a uma parte do domínio.
    - Em auditoria, inspecione métodos públicos e pontos de escrita em banco.
    """

    model = TabelaItemLinha
    template_name = "licitacoes/tabela_item_confirm_delete.html"

    def dispatch(self, request, *args, **kwargs):
        """Método `TabelaItemLinhaDeleteView.dispatch` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`, `request`, `*args`, `**kwargs`.
        """

        self.sessao = get_object_or_404(SessaoTermo, pk=kwargs["sessao_pk"])
        self.item = get_object_or_404(ItemSessao, pk=kwargs["item_pk"], sessao=self.sessao)
        return super().dispatch(request, *args, **kwargs)

    def get_queryset(self):
        """Método `TabelaItemLinhaDeleteView.get_queryset` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`.
        """

        return TabelaItemLinha.objects.filter(item=self.item)

    def get_context_data(self, **kwargs):
        """Método `TabelaItemLinhaDeleteView.get_context_data` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`, `**kwargs`.
        """

        context = super().get_context_data(**kwargs)
        context["next_url"] = _safe_next_url(
            self.request,
            f'{reverse("licitacoes_termo_detail", kwargs={"pk": self.sessao.termo.pk})}#tabela-itens-{self.item.pk}',
        )
        return context

    def form_valid(self, form):
        """Método `TabelaItemLinhaDeleteView.form_valid` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`, `form`.
        """

        response = super().form_valid(form)
        _normalize_tabela_linhas(self.item)
        return response

    def get_success_url(self):
        """Método `TabelaItemLinhaDeleteView.get_success_url` no fluxo do app `licitacoes`.

        Objetivo pedagógico:
        - Executa uma etapa específica de validação, transformação ou orquestração.
        - Verifique entradas, regras condicionais, retorno e efeitos colaterais.
        - Parâmetros: `self`.
        """

        default_url = f'{reverse("licitacoes_termo_detail", kwargs={"pk": self.sessao.termo.pk})}#tabela-itens-{self.item.pk}'
        return _safe_next_url(self.request, default_url)
